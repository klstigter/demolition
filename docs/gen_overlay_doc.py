"""
gen_overlay_doc.py
Generates a Word document that explains the black overlay bar
implementation on project task bars in the DHTMLX Gantt control add-in.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

OUTPUT = os.path.join(os.path.dirname(__file__), "RequestBar_Overlay_Technical_Doc.docx")

# ── helpers ──────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_code_block(doc, code_lines):
    """Add a shaded code block (monospace, light-grey background)."""
    for line in code_lines:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        # shade the paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        if bold_first and i == 0:
            cell.paragraphs[0].runs[0].bold = True

def set_table_header(table, headers):
    hdr_row = table.rows[0]
    for cell, h in zip(hdr_row.cells, headers):
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3864")
        tcPr.append(shd)

# ── build document ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# Title
title = doc.add_heading("Request-DayPlanning Overlay Bar — Technical Documentation", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Project: Demolition Gantt Control Add-In    |    Date: {datetime.date.today().strftime('%d %B %Y')}").italic = True
doc.add_paragraph()  # spacer

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "1. Purpose", level=1)
add_para(doc,
    "When a DayPlanning record in Business Central has plan_status = \"Request\" and carries a "
    "placeholder_date, the Gantt chart must visually signal that a requested (not yet confirmed) "
    "work slot exists within the task bar. This is done by injecting a semi-transparent black "
    "overlay div (.bc-req-bar) directly into the Gantt timeline DOM, positioned so it aligns "
    "exactly with the placeholder_date column of the owning project task bar.")

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "2. Background — Why a Custom Overlay?", level=1)
add_para(doc,
    "DHTMLX Gantt does not provide a built-in layer API that reliably supports custom per-day "
    "markers inside task bars. The gantt.addTaskLayer() approach was evaluated but discarded "
    "because:")
items = [
    "It fires inside a virtual-DOM render cycle and can be called before gantt.$task_data is "
    "fully mounted, producing silent failures.",
    "The layer callback has no guaranteed ordering relative to the default task renderer, "
    "making z-index stacking unpredictable.",
    "During drag operations DHTMLX bypasses the full render cycle (it moves the task element "
    "via direct style transforms), so addTaskLayer callbacks never re-fire.",
]
for item in items:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

add_para(doc,
    "\nThe direct DOM injection approach — appending absolutely-positioned divs to gantt.$task_data "
    "after each render — is deterministic, drag-aware (via explicit re-calls on onTaskDrag / "
    "onAfterTaskDrag), and does not require any private DHTMLX APIs.")

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "3. Data Flow", level=1)
add_para(doc, "The overlay requires two data sources to be loaded before it can render:")

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
set_table_header(tbl, ["Source", "Loaded by", "Stored in"])
rows = [
    ("Gantt tasks (JobTask)", "LoadProjectData()", "gantt internal store"),
    ("DayPlanning records", "LoadDayPlanningData()", "window.dayPlanningsByTask (index)"),
]
for r in rows:
    add_table_row(tbl, r)
doc.add_paragraph()

add_para(doc, "window.dayPlanningsByTask index structure:")
add_code_block(doc, [
    "// Key  = jobNo + \"|\" + jobTaskNo  (matches gantt task.id exactly)",
    "// Value = array of DayPlanning objects from BC",
    "window.dayPlanningsByTask = {",
    "  \"PROJ001|10\": [",
    "    { id: \"DT-1\", resource_id: \"R01\", work_date: \"2026-05-10\",",
    "      plan_status: \"Request\", placeholder_date: \"2026-05-10\",",
    "      hours: 8, ... },",
    "    ...",
    "  ],",
    "  ...",
    "}",
])

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "4. _renderRequestBars() — Step-by-Step", level=1)
add_para(doc,
    "The function _renderRequestBars() in wrapper.js is responsible for painting the black bars. "
    "It is intentionally kept outside the BOOT closure so it can be called from event handlers "
    "at module scope.")

add_heading(doc, "4.1  Resolve the DOM container", level=2)
add_para(doc,
    "The absolute-positioned task bars inside the Gantt timeline all live inside a single "
    "container div referenced by gantt.$task_data. In multi-view layouts (with a resource panel) "
    "this property may not exist at the root level, so a fallback via the UI view API is used:")
add_code_block(doc, [
    "var container = gantt.$task_data",
    "  || (gantt.$ui",
    "      && gantt.$ui.getView(\"timeline\")",
    "      && gantt.$ui.getView(\"timeline\").$task_data);",
    "if (!container) return;  // gantt not yet initialised",
])

add_heading(doc, "4.2  Remove stale bars from the previous render", level=2)
add_para(doc,
    "Because bars are appended as plain DOM children (not tracked by any virtual-DOM layer), "
    "they must be removed before each re-draw to avoid accumulation:")
add_code_block(doc, [
    "var old = container.querySelectorAll(\".bc-req-bar\");",
    "for (var i = 0; i < old.length; i++) {",
    "  old[i].parentNode && old[i].parentNode.removeChild(old[i]);",
    "}",
])

add_heading(doc, "4.3  Iterate gantt tasks and find Request DayPlannings", level=2)
add_para(doc,
    "gantt.eachTask() walks every loaded task. For each task the corresponding DayPlanning array "
    "is looked up from window.DayPlanningsByTask using the task's id (which equals jobNo|jobTaskNo).")
add_code_block(doc, [
    "gantt.eachTask(function (task) {",
    "  var dtList = window.DayPlanningsByTask[task.id];",
    "  if (!dtList || !dtList.length) return;  // nothing to render",
    "",
    "  // Only Request-status items with a placeholder_date get a bar",
    "  var seenDate = Object.create(null);      // deduplicate by date",
    "  for (var i = 0; i < dtList.length; i++) {",
    "    var dt = dtList[i];",
    "    if (dt.plan_status !== \"Request\" || !dt.placeholder_date) continue;",
    "    if (seenDate[dt.placeholder_date]) continue;  // one bar per date",
    "    seenDate[dt.placeholder_date] = true;",
    "    // ... create bar element (see 4.4)",
    "  }",
    "});",
])

add_heading(doc, "4.4  Calculate pixel position with getTaskPosition()", level=2)
add_para(doc,
    "gantt.getTaskPosition(task, fromDate, toDate) converts a date range into pixel coordinates "
    "relative to gantt.$task_data. It returns an object { left, top, height, width }.")
add_para(doc,
    "Two calls are made per bar:")
items2 = [
    "taskPos = getTaskPosition(task, task.start_date, task.end_date)  → gives the task row's "
    "top and height (used so the overlay bar covers the full task row height).",
    "pos = getTaskPosition(task, pd, pdEnd)  → pd is the placeholder_date parsed to midnight; "
    "pdEnd is pd + 1 day.  This gives the left offset and width of the single-day column.",
]
for item in items2:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

add_code_block(doc, [
    "var taskPos = gantt.getTaskPosition(task, task.start_date, task.end_date);",
    "",
    "var pd    = gantt.date.parseDate(dt.placeholder_date, \"%Y-%m-%d\");",
    "var pdEnd = gantt.date.add(new Date(pd.valueOf()), 1, \"day\");",
    "var pos   = gantt.getTaskPosition(task, pd, pdEnd);",
])

add_heading(doc, "4.5  Create and append the bar element", level=2)
add_para(doc, "A <div> is created and styled as an absolutely-positioned overlay:")
add_code_block(doc, [
    "var el = document.createElement(\"div\");",
    "el.className = \"bc-req-bar\";",
    "",
    "el.style.position     = \"absolute\";",
    "el.style.left         = pos.left + \"px\";       // X: placeholder_date column",
    "el.style.top          = taskPos.top + \"px\";    // Y: same row as task bar",
    "el.style.width        = Math.max(pos.width, 6) + \"px\";  // min 6px so it's always visible",
    "el.style.height       = taskPos.height + \"px\"; // full task-row height",
    "el.style.background   = \"rgba(0,0,0,0.85)\";   // dark semi-transparent black",
    "el.style.borderRadius = \"3px\";",
    "el.style.pointerEvents = \"none\";               // click-through — do not block task drag",
    "el.style.zIndex       = \"10\";                  // above task bar, below tooltips",
    "",
    "container.appendChild(el);",
])

add_para(doc,
    "The className \"bc-req-bar\" is the selector used for cleanup in step 4.2. "
    "pointerEvents:none is critical — without it the overlay would swallow mouse events "
    "and break Gantt drag-and-drop.")

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "5. Render Triggers", level=1)
add_para(doc,
    "The function must be called whenever the visual position of any task bar changes. "
    "Three triggers cover all cases:")

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
set_table_header(tbl2, ["Trigger", "Event / Hook", "Why needed"])
rows2 = [
    ("Full Gantt render",
     "gantt.attachEvent(\"onGanttRender\", ...)",
     "Fires after gantt.render() — data load, column toggle, layout change."),
    ("Real-time drag",
     "gantt.attachEvent(\"onTaskDrag\", ...)",
     "DHTMLX moves task bars via CSS transforms during drag WITHOUT calling gantt.render(). "
     "onTaskDrag fires on every mousemove tick, so bars track the task bar live."),
    ("Drag end",
     "gantt.attachEvent(\"onAfterTaskDrag\", ...)",
     "Snaps bars to the final resting position after the drag releases and "
     "auto-scheduling cascades have settled."),
]
for r in rows2:
    add_table_row(tbl2, r)
doc.add_paragraph()

add_code_block(doc, [
    "// 1. After every full render",
    "gantt.attachEvent(\"onGanttRender\", function () {",
    "  _updateResourceHeaderTooltip();",
    "  _renderRequestBars();",
    "});",
    "",
    "// 2. During drag — real-time tracking",
    "gantt.attachEvent(\"onTaskDrag\", function (id, mode, task, original) {",
    "  _renderRequestBars();",
    "  return true;",
    "});",
    "",
    "// 3. Drag end — final snap",
    "gantt.attachEvent(\"onAfterTaskDrag\", function (id, mode, e) {",
    "  _dragInProgress = true;",
    "  setTimeout(function () { _dragInProgress = false; }, 300);",
    "  _renderRequestBars();   // <-- snap bars to final position",
    "  // ... send OnJobTaskUpdated to BC",
    "  return true;",
    "});",
])

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "6. Black Bottom Strip in the Resource Panel", level=1)
add_para(doc,
    "The resource panel (lower half of the Gantt) shows per-resource workload cells "
    "(gantt_resource_marker). When a planned DayPlanning cell and a Request DayPlanning share the "
    "same date for the same resource, a thin black strip is appended inside the planned cell "
    "rather than painting a standalone grey box. This is handled inside renderResourceLine():")
add_code_block(doc, [
    "// Inside the timetable loop (planned cells):",
    "if (requestDateMap[dayStr]) {",
    "  var blackBar = document.createElement(\"div\");",
    "  blackBar.style.cssText = [",
    "    \"position:absolute\",",
    "    \"bottom:0\", \"left:0\", \"right:0\",",
    "    \"height:6px\",",
    "    \"background:#000000\",",
    "    \"border-radius:0 0 3px 3px\",",
    "    \"pointer-events:none\"",
    "  ].join(\";\");",
    "  cell.appendChild(blackBar);",
    "  requestDateHandled[dayStr] = true;",
    "}",
])
add_para(doc,
    "For dates where only a Request DayPlanning exists (no planned cell), a standalone grey "
    "marker (gantt_resource_marker_request, CSS #909090) is rendered instead.")

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "7. CSS Classes Summary", level=1)
tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = "Table Grid"
set_table_header(tbl3, ["Class / Selector", "Applied to", "Description"])
rows3 = [
    (".bc-req-bar",
     "Main timeline overlay div",
     "Black semi-transparent bar (rgba 0,0,0,0.85) inside project task bar row, "
     "one bar per unique placeholder_date."),
    (".gantt_resource_marker_request",
     "Resource panel marker cell",
     "Grey (#909090) standalone cell for dates with only Request tasks and no planned work."),
    ("inline blackBar style",
     "Child div inside planned marker cell",
     "6px black strip at the bottom of a green/red workload cell when a Request task "
     "shares the same date."),
]
for r in rows3:
    add_table_row(tbl3, r)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "8. Key Design Decisions", level=1)
decisions = [
    ("pointerEvents: none on overlay",
     "The .bc-req-bar must never intercept mouse events. Without this, hovering over "
     "the overlay would prevent DHTMLX from receiving mousedown/mousemove needed for "
     "task drag, effectively disabling drag on any date that has a Request bar."),
    ("Deduplication by placeholder_date (seenDate map)",
     "Multiple Request DayPlannings for different resources can share the same placeholder_date "
     "on the same task. Without deduplication, stacked identical bars would darken "
     "the overlay and produce visual artefacts."),
    ("Math.max(pos.width, 6) minimum width",
     "At very wide zoom levels or short column widths, getTaskPosition can return a "
     "sub-pixel width. A 6px minimum ensures the bar is always visible."),
    ("Class name .bc-req-bar as cleanup selector",
     "Using a unique project-namespaced class (bc- prefix) makes the querySelectorAll "
     "cleanup fast and avoids accidentally removing any DHTMLX-internal elements."),
    ("onTaskDrag fires on every mousemove",
     "This is intentional. Each call to _renderRequestBars() removes all .bc-req-bar "
     "elements and redraws them at the updated position. Because the function has no "
     "network I/O and operates only on DOM + in-memory data, the cost per frame is "
     "negligible for typical project sizes (<500 tasks)."),
]
for title_txt, body_txt in decisions:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(title_txt + ": ")
    run.bold = True
    p.add_run(body_txt)

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "9. Maintenance Notes", level=1)
notes = [
    "If the Gantt layout is changed from a single-view to a multi-view layout (or vice "
    "versa), verify that gantt.$task_data still resolves correctly. The fallback via "
    "gantt.$ui.getView(\"timeline\").$task_data handles the multi-view case but must be "
    "tested after any RecreateGanttLayout() change.",
    "If DHTMLX Gantt is upgraded, confirm that gantt.getTaskPosition() still accepts a "
    "task object (not just a task ID) and that $task_data remains a live DOM element reference.",
    "The window.DayPlanningsByTask index is rebuilt entirely on every LoadDayPlanningsData() call. "
    "If incremental DayPlanning updates (UpsertDayPlanning / DeleteDayPlanning) are used, verify that "
    "the index is also updated incrementally, otherwise _renderRequestBars() will use stale data.",
    "For large projects, consider throttling the onTaskDrag handler (e.g., requestAnimationFrame "
    "guard) if performance degradation is observed during drag on slow devices.",
]
for note in notes:
    p = doc.add_paragraph(style="List Number")
    p.add_run(note)

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "10. File Reference", level=1)
tbl4 = doc.add_table(rows=1, cols=2)
tbl4.style = "Table Grid"
set_table_header(tbl4, ["Item", "Location"])
refs = [
    ("Main implementation file", "src/dhx/ganttdemo2/wrapper.js"),
    ("_renderRequestBars() function", "wrapper.js — after BOOT() closure, ~line 1068"),
    ("Render triggers (onGanttRender, onTaskDrag, onAfterTaskDrag)", "wrapper.js — inside BOOT(), after gantt.init()"),
    ("renderResourceLine() — resource panel strip", "wrapper.js — calculateResourceLoad() block"),
    ("DayPlanning data loading", "wrapper.js — LoadDayPlanningsData()"),
]
for r in refs:
    add_table_row(tbl4, r)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("— End of Document —").italic = True

# ── save ─────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Document saved: {OUTPUT}")
