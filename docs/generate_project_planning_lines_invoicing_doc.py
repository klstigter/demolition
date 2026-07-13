"""
generate_project_planning_lines_invoicing_doc.py
Generates the technical documentation for "Project Planning Lines for Invoicing"
(the Day-Planning-to-Invoice Release 1 feature) as a Word (.docx) file using python-docx.

Run:  python generate_project_planning_lines_invoicing_doc.py
Output: Project_Planning_Lines_for_Invoicing.docx  (same folder as this script)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ─── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def header_row(table, *texts, bg='1F3964'):
    row = table.rows[0]
    for i, text in enumerate(texts):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)


def data_row(table, *values, shade=False):
    row = table.add_row()
    for i, val in enumerate(values):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        if shade:
            set_cell_bg(cell, 'EBF3FB')
        p = cell.paragraphs[0]
        p.add_run(str(val)).font.size = Pt(10)
    return row


def add_table(doc, col_widths_cm, *header_texts, bg='1F3964'):
    cols = len(header_texts)
    tbl = doc.add_table(rows=1, cols=cols)
    tbl.style = 'Table Grid'
    header_row(tbl, *header_texts, bg=bg)
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[j].width = Cm(w)
    return tbl


def add_code(doc, code: str):
    """Monospace code block with light-blue background."""
    for line in code.split('\n'):
        para = doc.add_paragraph()
        para.style = doc.styles['No Spacing']
        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EDF2F8')
        pPr.append(shd)
    doc.add_paragraph()


def info_box(doc, text: str, label='Note', color='D6E4F7'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    run = p.add_run(f'{label}\n{text}')
    run.font.size = Pt(10)
    doc.add_paragraph()


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text=''):
    return doc.add_paragraph(text)


def bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


# ════════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Project Planning Lines for Invoicing')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run('DailyOptimizer Extension - Day-Planning-to-Invoice, Release 1 - Technical Documentation').font.size = Pt(14)

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(
    f'Extension: DailyOptimizer  |  Publisher: Optimizers  |  Version: 28.0.0.3\n'
    f'BC Runtime: 15.0  |  Generated: {datetime.datetime.now().strftime("%d %B %Y %H:%M")}'
).font.size = Pt(10)

doc.add_page_break()

# ── 1. OVERVIEW ───────────────────────────────────────────────────────────────

h(doc, '1. Overview')
p(doc,
  'This document covers the full implementation of "Project Planning Lines for Invoicing" '
  '(the Day-Planning-to-Invoice Release 1 feature) in the DailyOptimizer AL extension for '
  'Microsoft Dynamics 365 Business Central. It implements the design in '
  '"Day_Planning2Invoice.docx" (Release 1 scope): realized Day Planning usage is posted to '
  'standard BC Job Ledger Entries (pre-existing traceability, unchanged by this feature), '
  'then summarized - grouped by Skill - into standard Project Planning Lines (Job Planning '
  'Line) ready for the native BC invoice flow, while a lean link table preserves a '
  'one-to-many trace from each summary line back to the individual posted usage entries '
  'behind it.')
p(doc,
  'The commercial invoice therefore stays summarized per Skill/Invoice Resource, while the '
  'customer-facing detail report can still show the real per-resource hours behind each '
  'summary line. Standard BC invoice creation (Project Planning Lines -> Sales Invoice -> '
  'Posted Sales Invoice) is used unchanged from the generated Job Planning Lines onward - '
  'this feature only bridges "posted Day Planning usage" to "billable Job Planning Line".')

h(doc, '1.1  End-to-End Flow', 2)
tbl = add_table(doc, [1.5, 15.5], 'Step', 'Description')
flow_steps = [
    ('1', 'Day Planning line is realized/posted -> standard BC Job Ledger Entry is created '
          '(and, for Resource lines, a Res. Ledger Entry). Pre-existing behaviour, not part '
          'of this feature - see codeunit 50603 "EventSubs" section below for how the link '
          'back to the originating Day Planning row is recorded at that moment.'),
    ('2', 'User runs "Prepare Invoice Lines" (or its batch-report / multi-select variants) '
          'from the Day Plannings list or the Work Order card.'),
    ('3', 'Codeunit 50607 "Job Invoice Prep. Mgt." selects posted, not-yet-linked Day '
          'Planning rows, resolves each one''s Skill -> "Invoice Resource No.", and groups '
          'candidates by Job No. + Job Task No. + Invoice Resource No. + Unit of Measure.'),
    ('4', 'One billable Job Planning Line is created per group (Type = Resource, Line Type '
          '= Billable), and one "Job Ledger Invoice Link" row is written per included usage '
          'entry, pointing at that summary line.'),
    ('5', 'Standard BC invoice flow continues unchanged: Project Planning Lines -> Sales '
          'Invoice -> Posted Sales Invoice.'),
    ('6', 'Report 50606 "Job Planning Line Usage Detail" (or the drilldown on page 50681) '
          'lets anyone trace a summary Job Planning Line back to the individual posted '
          'usage entries/resources behind it, for customer-facing detail or support.'),
]
for i, row in enumerate(flow_steps):
    data_row(tbl, *row, shade=(i % 2 == 0))
doc.add_paragraph()

info_box(doc,
    'Grouping key = Job No. + Job Task No. + Skill''s Invoice Resource No. + Unit of '
    'Measure Code - NOT Unit Price. Unit Price is resolved once per group via standard '
    'Job Planning Line field-validation (resource price list / job price group / work '
    'type), exactly as a manual planning-line entry would resolve it. This is the one '
    'open design decision the original design doc explicitly left to the programmer.',
    label='Key Design Decision', color='FFF2CC')

# ── 2. OBJECT INVENTORY ───────────────────────────────────────────────────────

h(doc, '2. Complete Object Inventory')
p(doc, 'The following AL objects were created or modified as part of this feature:')

tbl = add_table(doc, [2.3, 1.7, 5.2, 2.3, 5.5],
                'Object Type', 'Object ID', 'Object Name', 'Status', 'Change Summary')
rows = [
    ('Table',           '50614', 'Job Ledger Invoice Link',        'NEW',      'Traceability link: posted usage entry -> summary Job Planning Line'),
    ('Table Extension', '50609', 'Opt. Skill Code',                'NEW',      'Adds "Invoice Resource No." to standard Skill Code table'),
    ('Codeunit',        '50607', 'Job Invoice Prep. Mgt.',         'NEW',      'Core business logic: selects, groups, and creates billable Job Planning Lines'),
    ('Codeunit',        '50603', 'EventSubs',                      'Modified', 'Adds delete-cleanup subscriber on Job Planning Line OnAfterDeleteEvent'),
    ('Report',          '50606', 'Job Planning Line Usage Detail', 'NEW',      'Customer detail hours report - traces a summary line to its usage entries'),
    ('Report',          '50607', 'Prepare Proj. Planning Lines',   'NEW',      'Batch-report entry point with request-page filtering'),
    ('Report',          '50600', 'RepairData',                     'Modified', 'Added DeleteGeneratedInvoicePlanningLines repair action'),
    ('Page',            '50630', 'Day Plannings',                  'Modified', 'Added "Invoicing" action group with 3 actions (see Section 6.1)'),
    ('Page',            '50662', 'Workorder Card',                 'Modified', 'Added "Prepare Invoice Lines" action'),
    ('Page',             '50681', 'Job Ledger Invoice Link',        'NEW',      'Read-only traceability list, with drilldown to the summary line'),
    ('Page Extension',  '50626', 'Skill Codes Opt.',               'NEW',      'Adds "Invoice Resource No." field to the standard Skill Codes page'),
    ('Page',            '50612', 'Planning Role Center',           'Modified', 'Added "Skill Codes" and "Job Ledger Invoice Link" nav actions'),
]
for i, r in enumerate(rows):
    dr = data_row(tbl, *r, shade=(i % 2 == 0))
    if r[3] == 'NEW':
        set_cell_bg(dr.cells[3], 'C6EFCE')
doc.add_paragraph()

info_box(doc,
    'Report 50607 (a Report object) and Codeunit 50607 (a Codeunit object) intentionally '
    'share the numeric ID 50607 - AL object IDs are namespaced per object TYPE, so this is '
    'not a collision. Likewise Page 50612 "Planning Role Center" is a pre-existing object '
    'only modified by this feature, unrelated to Codeunit/Report 50612 (which do not exist).',
    label='ID Namespace Note', color='D6E4F7')

# ── 3. DATA MODEL ─────────────────────────────────────────────────────────────

h(doc, '3. Data Model')

h(doc, '3.1  Table 50614 "Job Ledger Invoice Link" [NEW]', 2)
p(doc,
  'A small, intentionally non-batch companion table. One row per posted Job Ledger Entry '
  '(resource usage) that has been rolled into a billable Job Planning Line. It exists purely '
  'for performance, anti-double-invoicing, and drilldown from a summary line back to its '
  'detailed usage entries - it does not replace or duplicate the Job Ledger Entry itself.')

tbl = add_table(doc, [1.3, 4.5, 2.2, 2.0, 8.0],
                'Field No.', 'Field Name', 'Type', 'Editable', 'Purpose')
field_rows = [
    ('1', 'Job Ledger Entry No.', 'Integer', 'No',
     'Primary Key (Clustered). TableRelation to "Job Ledger Entry"."Entry No.". The PK on '
     'this field ALONE is what prevents double-invoicing: a usage entry can only ever be '
     'linked to one invoice planning line, ever.'),
    ('2', 'Job No.', 'Code[20]', 'No', 'TableRelation to Job. Denormalized from the Job Planning Line for direct filtering/reporting.'),
    ('3', 'Job Task No.', 'Code[20]', 'No', 'TableRelation to Job Task (scoped by Job No.).'),
    ('4', 'Invoice Job Planning Line No.', 'Integer', 'No', 'The "Line No." of the billable Job Planning Line this usage entry was rolled into.'),
    ('5', 'Skill Code', 'Code[20]', 'No',
     'Optional, for reporting only - the Skill this specific usage entry originated from '
     '(via its Day Planning row). Kept per-entry rather than per invoice line, because two '
     'different Skills can share the same Invoice Resource No.'),
]
for i, r in enumerate(field_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '3.1.1  Keys', 3)
tbl = add_table(doc, [4, 6, 8], 'Key', 'Fields', 'Purpose')
key_rows = [
    ('PK (Clustered)', 'Job Ledger Entry No.', 'Anti-double-invoicing guard (see above).'),
    ('SummaryLookup', 'Job No., Job Task No., Invoice Job Planning Line No.',
     'Reverse lookup from a Job Planning Line back to all its linked usage entries - used '
     'by Report 50606 and the delete-cleanup subscriber.'),
]
for i, r in enumerate(key_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '3.1.2  Full AL Source', 3)
add_code(doc,
r"""table 50614 "Job Ledger Invoice Link"
{
    // The Primary Key on "Job Ledger Entry No." alone is what prevents double-invoicing:
    // a usage entry can only ever be linked to one invoice planning line, ever. Do not
    // weaken this (e.g. do not add a composite key that would allow a second row for the
    // same Job Ledger Entry No.).
    DataClassification = CustomerContent;
    Caption = 'Job Ledger Invoice Link';

    fields
    {
        field(1; "Job Ledger Entry No."; Integer)
        {
            Caption = 'Job Ledger Entry No.';
            DataClassification = CustomerContent;
            TableRelation = "Job Ledger Entry";
            Editable = false;
        }
        field(2; "Job No."; Code[20])
        {
            Caption = 'Job No.';
            DataClassification = CustomerContent;
            TableRelation = Job;
            Editable = false;
        }
        field(3; "Job Task No."; Code[20])
        {
            Caption = 'Job Task No.';
            DataClassification = CustomerContent;
            TableRelation = "Job Task"."Job Task No." where("Job No." = field("Job No."));
            Editable = false;
        }
        field(4; "Invoice Job Planning Line No."; Integer)
        {
            Caption = 'Invoice Job Planning Line No.';
            DataClassification = CustomerContent;
            Editable = false;
        }
        field(5; "Skill Code"; Code[20])
        {
            Caption = 'Skill Code';
            DataClassification = CustomerContent;
            TableRelation = "Skill Code";
            Editable = false;
        }
    }

    keys
    {
        key(PK; "Job Ledger Entry No.")
        {
            Clustered = true;
        }
        key(SummaryLookup; "Job No.", "Job Task No.", "Invoice Job Planning Line No.")
        {
        }
    }

    fieldgroups
    {
        fieldgroup(DropDown; "Job No.", "Job Task No.", "Invoice Job Planning Line No.")
        {
        }
    }
}""")

h(doc, '3.2  Table Extension 50609 "Opt. Skill Code" [NEW]', 2)
p(doc,
  'Extends the standard "Skill Code" table with the Resource that billable Job Planning '
  'Lines should be created against when summarizing posted usage for that Skill. Invoice '
  'preparation hard-stops with a clear error if this is blank for any Skill actually in use '
  'on posted, unlinked Job Ledger Entries - per the design doc''s explicit safeguard.')
add_code(doc,
r"""tableextension 50609 "Opt. Skill Code" extends "Skill Code"
{
    fields
    {
        field(50609; "Invoice Resource No."; Code[20])
        {
            Caption = 'Invoice Resource No.';
            DataClassification = CustomerContent;
            TableRelation = Resource."No.";
            ToolTip = 'Specifies the resource that billable planning lines for this skill '
                      + 'are created against when preparing invoice lines from posted Day '
                      + 'Planning usage.';
        }
    }
}""")
p(doc, 'Exposed for data entry via Page Extension 50626 "Skill Codes Opt." - see Section 6.4.')

# ── 4. CORE BUSINESS LOGIC ────────────────────────────────────────────────────

h(doc, '4. Codeunit 50607 "Job Invoice Prep. Mgt." [NEW]')
p(doc,
  'The single source of truth for turning posted, not-yet-invoiced Day Planning usage into '
  'billable Job Planning Lines. All UI entry points (page actions, the batch report) call '
  'into this codeunit - none of them re-implement the selection/grouping logic.')

h(doc, '4.1  Design Principles', 2)
bullet(doc, 'Driven from Day Planning, not Job Ledger Entry. "Day Planning".Posted = true is the '
             'authoritative signal that a row has been posted into a Job Ledger Entry, and '
             '"Day Planning"."Job Entry No." (field 151) is where codeunit 50603 recorded that '
             'Job Ledger Entry''s Entry No. at posting time. The candidate loop therefore reads '
             '"Job Entry No." directly - including into the link table''s own "Job Ledger Entry '
             'No." field - rather than looping Job Ledger Entry independently and relying on an '
             'unenforced coincidental equality between the two.')
bullet(doc, 'The matching Job Ledger Entry is still fetched via Get(), purely to read Quantity/'
             'Unit of Measure and to scope the feature to resource usage (Entry Type = Usage, '
             'Type = Resource).')
bullet(doc, 'Two-pass processing. Pass 1 resolves every candidate''s Skill -> Invoice Resource '
             'No. and builds the grouping with NO database writes at all, so a missing Invoice '
             'Resource No. on any Skill in play raises Error() before anything is written. Pass '
             '2 then creates the Job Planning Lines and Job Ledger Invoice Link rows. Since no '
             'COMMIT() is issued anywhere in this codeunit, a runtime error during Pass 2 rolls '
             'back the whole run via BC''s ambient transaction - all-or-nothing, with no '
             'explicit transaction wrapper needed.')
bullet(doc, 'Per-group entry lists are carried as a Dictionary of [Text, List of [Integer]] - '
             'each group''s Job Entry No. values stay typed Integers end-to-end, with no Format/'
             'Split/Evaluate round-trip through Text (an earlier revision used a comma-joined '
             'Text approach; it was replaced because Evaluate()''s success was never checked, '
             'risking silent data corruption on a parse failure). Every touch of a group''s list '
             'follows an explicit Get-mutate-Set-back idiom - Get() the List into a local '
             'variable, mutate the local copy, then Set() it back into the dictionary - rather '
             'than assuming in-place mutation persists, which is safe regardless of List''s '
             'value/reference assignment semantics in a given AL version.')
bullet(doc, 'Only Day Planning rows that have actually been posted (Posted = true, "Job Entry '
             'No." <> 0) are candidates. If the Job Ledger Entry they point to can no longer be '
             'found (a data-integrity edge case that should not normally happen), that row is '
             'silently skipped rather than erroring the whole batch.')

h(doc, '4.2  Public Procedures', 2)
tbl = add_table(doc, [7, 10], 'Procedure', 'Purpose')
proc_rows = [
    ('PrepareInvoiceLines(JobNo, JobTaskNo, var 4 counters): LinesCreated',
     'Sweeps EVERY posted, unlinked Day Planning row for the given Job (optionally scoped '
     'to one Job Task). Used by the Work Order card - "invoice everything ready for this '
     'job task".'),
    ('PrepareInvoiceLinesForSelection(var SelectedDayPlanning, var 4 counters): LinesCreated',
     'Operates ONLY on the caller-supplied, already-filtered Day Planning recordset - e.g. '
     'a user''s exact multi-selection on the Day Plannings list, or a batch report''s '
     'request-page-filtered dataset. Does not sweep beyond what the caller already scoped.'),
    ('[TryFunction] TryPrepareInvoiceLines(...)', 'Try-wrapper around PrepareInvoiceLines, for batch callers processing several Job/Job Task pairs that want one failure isolated rather than aborting the whole run.'),
    ('[TryFunction] TryPrepareInvoiceLinesForSelection(...)', 'Try-wrapper around PrepareInvoiceLinesForSelection, same isolation semantics.'),
    ('FormatResultMessage(LinesCreated, 4 counters): Text',
     'Builds the single, shared human-readable breakdown message used by every caller (see '
     'Section 4.3), so the wording is defined exactly once.'),
]
for i, r in enumerate(proc_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

info_box(doc,
    'AL''s TryFunction mechanism rolls back DATABASE changes on failure, but NOT local/'
    'var-parameter variable state accumulated before the runtime error. So even when Pass 1 '
    'Error()s partway through (e.g. a Skill with no Invoice Resource No.), the four '
    'breakdown counters accumulated up to that point remain meaningful and are still shown '
    'to the user alongside the error text.',
    label='TryFunction Semantics', color='D6E4F7')

h(doc, '4.3  The Breakdown-Counter Message', 2)
p(doc,
  'A bare "N Project Planning Line(s) created." message gives no insight into WHY nothing '
  'happened when N = 0. FormatResultMessage instead reports four buckets - only the non-zero '
  'ones are shown, so the clean-success case stays uncluttered:')
tbl = add_table(doc, [4, 13], 'Counter', 'Meaning')
counter_rows = [
    ('ProcessedCount', 'Rows successfully grouped into a Project Planning Line.'),
    ('AlreadyLinkedCount', 'Rows skipped because a Job Ledger Invoice Link already exists for them.'),
    ('NotPostedCount', 'Rows skipped because they are not (yet) posted to a Job Ledger Entry.'),
    ('SkippedOtherCount', 'Rows skipped for other reasons (Job Ledger Entry not found, or not Usage/Resource type) - rare, a data-integrity edge case.'),
]
for i, r in enumerate(counter_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()
p(doc, 'Example output: "6 record(s) processed into 2 Project Planning Line(s).\\n'
       '8 record(s) already in a Project Planning Line.\\n'
       '25 record(s) have not been posted into Job Ledger Entries yet."')

h(doc, '4.4  Full AL Source', 2)
add_code(doc,
r"""codeunit 50607 "Job Invoice Prep. Mgt."
{
    Permissions = tabledata "Job Ledger Entry" = r,
                  tabledata "Day Planning" = r,
                  tabledata "Skill Code" = r,
                  tabledata "Job Planning Line" = rim,
                  tabledata "Job Ledger Invoice Link" = ri;

    trigger OnRun()
    begin
    end;

    procedure PrepareInvoiceLines(JobNo: Code[20]; JobTaskNo: Code[20]; var ProcessedCount: Integer; var AlreadyLinkedCount: Integer; var NotPostedCount: Integer; var SkippedOtherCount: Integer) LinesCreated: Integer
    var
        DayPlanning: Record "Day Planning";
    begin
        DayPlanning.SetRange("Job No.", JobNo);
        if JobTaskNo <> '' then
            DayPlanning.SetRange("Job Task No.", JobTaskNo);
        DayPlanning.SetRange(Posted, true);
        exit(PrepareInvoiceLinesFromDayPlanning(DayPlanning, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount));
    end;

    procedure PrepareInvoiceLinesForSelection(var SelectedDayPlanning: Record "Day Planning"; var ProcessedCount: Integer; var AlreadyLinkedCount: Integer; var NotPostedCount: Integer; var SkippedOtherCount: Integer) LinesCreated: Integer
    begin
        exit(PrepareInvoiceLinesFromDayPlanning(SelectedDayPlanning, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount));
    end;

    local procedure PrepareInvoiceLinesFromDayPlanning(var DayPlanning: Record "Day Planning"; var ProcessedCount: Integer; var AlreadyLinkedCount: Integer; var NotPostedCount: Integer; var SkippedOtherCount: Integer) LinesCreated: Integer
    var
        JobLedgerEntry: Record "Job Ledger Entry";
        JobLedgerInvoiceLink: Record "Job Ledger Invoice Link";
        SkillCodeRec: Record "Skill Code";
        GroupHours: Dictionary of [Text, Decimal];
        GroupJobNo: Dictionary of [Text, Code[20]];
        GroupJobTaskNo: Dictionary of [Text, Code[20]];
        GroupInvoiceResNo: Dictionary of [Text, Code[20]];
        GroupUOM: Dictionary of [Text, Code[10]];
        GroupEntryNos: Dictionary of [Text, List of [Integer]];
        EntrySkillCode: Dictionary of [Integer, Code[20]];
        GroupKeys: List of [Text];
        TempEmptyIntList: List of [Integer];
        TempIntList: List of [Integer];
        GroupKey: Text;
        SkillCode: Code[20];
        InvoiceResNo: Code[20];
    begin
        if DayPlanning.FindSet() then
            repeat
                if DayPlanning.Posted and (DayPlanning."Job Entry No." <> 0) then begin
                    if JobLedgerInvoiceLink.Get(DayPlanning."Job Entry No.") then
                        AlreadyLinkedCount += 1
                    else
                        if JobLedgerEntry.Get(DayPlanning."Job Entry No.") then begin
                            if (JobLedgerEntry."Entry Type" = JobLedgerEntry."Entry Type"::Usage) and
                               (JobLedgerEntry.Type = JobLedgerEntry.Type::Resource)
                            then begin
                                SkillCode := DayPlanning.Skill;
                                if SkillCode = '' then
                                    Error(NoSkillOnDayPlanningErr, DayPlanning."Job No.", DayPlanning."Job Task No.", DayPlanning."Day Line No.");

                                if not SkillCodeRec.Get(SkillCode) then
                                    Error(SkillCodeNotFoundErr, SkillCode);

                                InvoiceResNo := SkillCodeRec."Invoice Resource No.";
                                if InvoiceResNo = '' then
                                    Error(NoInvoiceResourceErr, SkillCode);

                                GroupKey := StrSubstNo('%1|%2|%3|%4',
                                    DayPlanning."Job No.", DayPlanning."Job Task No.",
                                    InvoiceResNo, JobLedgerEntry."Unit of Measure Code");

                                if not GroupHours.ContainsKey(GroupKey) then begin
                                    GroupHours.Add(GroupKey, 0);
                                    GroupJobNo.Add(GroupKey, DayPlanning."Job No.");
                                    GroupJobTaskNo.Add(GroupKey, DayPlanning."Job Task No.");
                                    GroupInvoiceResNo.Add(GroupKey, InvoiceResNo);
                                    GroupUOM.Add(GroupKey, JobLedgerEntry."Unit of Measure Code");
                                    Clear(TempEmptyIntList);
                                    GroupEntryNos.Add(GroupKey, TempEmptyIntList);
                                    GroupKeys.Add(GroupKey);
                                end;

                                GroupHours.Set(GroupKey, GroupHours.Get(GroupKey) + JobLedgerEntry.Quantity);
                                TempIntList := GroupEntryNos.Get(GroupKey);
                                TempIntList.Add(DayPlanning."Job Entry No.");
                                GroupEntryNos.Set(GroupKey, TempIntList);

                                if not EntrySkillCode.ContainsKey(DayPlanning."Job Entry No.") then
                                    EntrySkillCode.Add(DayPlanning."Job Entry No.", SkillCode);

                                ProcessedCount += 1;
                            end else
                                SkippedOtherCount += 1;
                        end else
                            SkippedOtherCount += 1;
                end else
                    NotPostedCount += 1;
            until DayPlanning.Next() = 0;

        foreach GroupKey in GroupKeys do begin
            CreateInvoicePlanningLine(
                GroupJobNo.Get(GroupKey), GroupJobTaskNo.Get(GroupKey),
                GroupInvoiceResNo.Get(GroupKey), GroupUOM.Get(GroupKey), GroupHours.Get(GroupKey),
                GroupEntryNos.Get(GroupKey), EntrySkillCode);
            LinesCreated += 1;
        end;
        exit(LinesCreated);
    end;

    local procedure CreateInvoicePlanningLine(JobNo: Code[20]; JobTaskNo: Code[20]; InvoiceResNo: Code[20]; UOMCode: Code[10]; Hours: Decimal; EntryNos: List of [Integer]; var EntrySkillCode: Dictionary of [Integer, Code[20]])
    var
        JobPlanningLine: Record "Job Planning Line";
        JobLedgerInvoiceLink: Record "Job Ledger Invoice Link";
        EntryNo: Integer;
    begin
        JobPlanningLine.Init();
        JobPlanningLine."Job No." := JobNo;
        JobPlanningLine."Job Task No." := JobTaskNo;
        JobPlanningLine."Line No." := GetNextJobPlanningLineNo(JobNo, JobTaskNo);
        JobPlanningLine.Insert(true);

        JobPlanningLine.Validate("Line Type", JobPlanningLine."Line Type"::Billable);
        JobPlanningLine.Validate(Type, JobPlanningLine.Type::Resource);
        JobPlanningLine.Validate("No.", InvoiceResNo);
        if JobPlanningLine."Unit of Measure Code" <> UOMCode then
            JobPlanningLine.Validate("Unit of Measure Code", UOMCode);
        JobPlanningLine.Validate(Quantity, Hours);
        JobPlanningLine.Modify(true);

        foreach EntryNo in EntryNos do begin
            JobLedgerInvoiceLink.Init();
            JobLedgerInvoiceLink."Job Ledger Entry No." := EntryNo;
            JobLedgerInvoiceLink."Job No." := JobPlanningLine."Job No.";
            JobLedgerInvoiceLink."Job Task No." := JobPlanningLine."Job Task No.";
            JobLedgerInvoiceLink."Invoice Job Planning Line No." := JobPlanningLine."Line No.";
            JobLedgerInvoiceLink."Skill Code" := EntrySkillCode.Get(EntryNo);
            JobLedgerInvoiceLink.Insert(true);
        end;
    end;

    local procedure GetNextJobPlanningLineNo(JobNo: Code[20]; JobTaskNo: Code[20]): Integer
    var
        JobPlanningLine: Record "Job Planning Line";
    begin
        JobPlanningLine.SetRange("Job No.", JobNo);
        JobPlanningLine.SetRange("Job Task No.", JobTaskNo);
        if JobPlanningLine.FindLast() then;
        exit(JobPlanningLine."Line No." + 10000);
    end;
}""")

# ── 5. TRACEABILITY & CLEANUP ─────────────────────────────────────────────────

h(doc, '5. Codeunit 50603 "EventSubs" - Traceability & Cleanup')
p(doc,
  'This codeunit already carried the Day-Planning-to-Job-Ledger-Entry traceability that '
  'this feature builds on (unchanged by this feature - see 5.1). One new subscriber was '
  'added: cleanup of Job Ledger Invoice Link rows when their generated Job Planning Line is '
  'deleted (5.2).')

h(doc, '5.1  Pre-Existing: Posting-Time Traceability (unchanged)', 2)
p(doc,
  'Two subscribers on Codeunit "Job Jnl.-Post Line" copy Day Planning traceability fields '
  'onto the Job Ledger Entry (and, for Resource lines, the Res. Ledger Entry) at posting '
  'time, and mark the originating Day Planning row Posted with its resulting entry numbers. '
  'This is what makes "Day Planning"."Job Entry No." (field 151) - the field this feature''s '
  'entire candidate-selection logic is built on - available in the first place.')

h(doc, '5.2  New: Delete-Cleanup Subscriber [Added]', 2)
p(doc,
  'When a generated Job Planning Line is deleted (BC''s own delete logic already allowed '
  'it - this subscriber reacts after the fact and never fights standard BC delete logic), '
  'the matching Job Ledger Invoice Link rows are dropped, so the underlying usage entries '
  'become invoiceable again on the next "Prepare Invoice Lines" run.')
add_code(doc,
r"""[EventSubscriber(ObjectType::Table, Database::"Job Planning Line", 'OnAfterDeleteEvent', '', false, false)]
local procedure JobPlanningLine_OnAfterDeleteEvent(var Rec: Record "Job Planning Line"; RunTrigger: Boolean)
var
    JobLedgerInvoiceLink: Record "Job Ledger Invoice Link";
begin
    if Rec.IsTemporary() then
        exit;

    JobLedgerInvoiceLink.SetRange("Job No.", Rec."Job No.");
    JobLedgerInvoiceLink.SetRange("Job Task No.", Rec."Job Task No.");
    JobLedgerInvoiceLink.SetRange("Invoice Job Planning Line No.", Rec."Line No.");
    JobLedgerInvoiceLink.DeleteAll(false);
end;""")

info_box(doc,
    'Deleting a Job Ledger Invoice Link row manually (bypassing the Job Planning Line '
    'delete) does NOT adjust the summary line''s quantity - it only re-opens the usage '
    'entry for re-invoicing, creating a double-billing risk if the original summary line''s '
    'hours are not also corrected. This is exactly why Page 50681''s DeleteAllowed is '
    'false - see Section 6.3.',
    label='Why Manual Link Deletion Is Blocked', color='FFF2CC')

# ── 6. USER INTERFACE ─────────────────────────────────────────────────────────

h(doc, '6. User Interface')

h(doc, '6.1  Page 50630 "Day Plannings" - "Invoicing" Action Group', 2)
p(doc, 'Three actions, all promoted to the ribbon under an "Invoicing" dropdown:')

tbl = add_table(doc, [4.5, 12.5], 'Action', 'Behaviour')
action_rows = [
    ('Prepare Invoice Lines',
     'User multi-selects Day Planning rows. CurrPage.SetSelectionFilter captures the '
     'exact selection, which is handed directly to PrepareInvoiceLinesForSelection - only '
     'the rows the user actually selected are considered, not a broad Job/Job Task sweep. '
     'Result shown via the shared FormatResultMessage breakdown.'),
    ('Prepare Project Planning Lines for Invoicing... (Batch-Report)',
     'RunObject = Report 50607 "Prepare Proj. Planning Lines" - opens a request page '
     '(filterable by Job No., Job Task No., Skill, Assigned Resource No.) then runs the '
     'same PrepareInvoiceLinesForSelection logic once against the filtered dataset. See '
     'Section 7.2.'),
    ('Project Planning Lines',
     'Navigation-only: for the selected Day Planning rows, resolves each row''s Job '
     'Entry No. through Job Ledger Invoice Link (table 50614) to find the Job No./Job '
     'Task No./Line No. it was rolled into, then opens the standard "Job Planning Lines" '
     'page filtered to the resolved set. See note below on filter precision.'),
]
for i, r in enumerate(action_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

info_box(doc,
    '"Project Planning Lines" builds its filter as three independent OR-lists (Job No., '
    'Job Task No., Line No.) ANDed together, because AL cannot express an arbitrary '
    'composite-tuple OR filter ("(A=1 AND B=2 AND C=3) OR (A=4 AND B=5 AND C=6)") as a '
    'single filter view. This is EXACT when the whole selection resolves to one Job No./'
    'Job Task No. pair (the common case), and may be a mild OVER-approximation - possibly '
    'including a few extra unrelated lines - when the selection spans multiple distinct '
    'pairs. Acceptable for a browse/navigate convenience action; not used anywhere that '
    'writes data.',
    label='Filter Precision Caveat', color='FFF2CC')

h(doc, '6.1.1  OnAction - "Prepare Invoice Lines"', 3)
add_code(doc,
r"""trigger OnAction()
var
    JobInvoicePrepMgt: Codeunit "Job Invoice Prep. Mgt.";
    SelectedDayPlanning: Record "Day Planning";
    LinesCreated, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount: Integer;
begin
    CurrPage.SetSelectionFilter(SelectedDayPlanning);
    if not SelectedDayPlanning.FindSet() then begin
        Message(NothingSelectedMsg);
        exit;
    end;

    if JobInvoicePrepMgt.TryPrepareInvoiceLinesForSelection(SelectedDayPlanning, LinesCreated, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount) then
        Message(JobInvoicePrepMgt.FormatResultMessage(LinesCreated, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount))
    else
        Message('%1\%2',
            JobInvoicePrepMgt.FormatResultMessage(LinesCreated, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount),
            StrSubstNo(FailedLbl, GetLastErrorText()));
end;""")

h(doc, '6.1.2  OnAction - "Project Planning Lines"', 3)
add_code(doc,
r"""trigger OnAction()
var
    SelectedDayPlanning: Record "Day Planning";
    JobLedgerInvoiceLink: Record "Job Ledger Invoice Link";
    JobPlanningLine: Record "Job Planning Line";
    JobPlanningLinesPage: Page "Job Planning Lines";
    JobNos: List of [Code[20]];
    JobTaskNos: List of [Code[20]];
    LineNos: List of [Integer];
begin
    CurrPage.SetSelectionFilter(SelectedDayPlanning);
    if not SelectedDayPlanning.FindSet() then begin
        Message(NothingSelectedMsg);
        exit;
    end;

    repeat
        if SelectedDayPlanning."Job Entry No." <> 0 then
            if JobLedgerInvoiceLink.Get(SelectedDayPlanning."Job Entry No.") then begin
                if not JobNos.Contains(JobLedgerInvoiceLink."Job No.") then
                    JobNos.Add(JobLedgerInvoiceLink."Job No.");
                if not JobTaskNos.Contains(JobLedgerInvoiceLink."Job Task No.") then
                    JobTaskNos.Add(JobLedgerInvoiceLink."Job Task No.");
                if not LineNos.Contains(JobLedgerInvoiceLink."Invoice Job Planning Line No.") then
                    LineNos.Add(JobLedgerInvoiceLink."Invoice Job Planning Line No.");
            end;
    until SelectedDayPlanning.Next() = 0;

    if JobNos.Count() = 0 then begin
        Message(NoLinksFoundMsg);
        exit;
    end;

    JobPlanningLine.SetFilter("Job No.", BuildCodeOrFilter(JobNos));
    JobPlanningLine.SetFilter("Job Task No.", BuildCodeOrFilter(JobTaskNos));
    JobPlanningLine.SetFilter("Line No.", BuildIntegerOrFilter(LineNos));
    JobPlanningLinesPage.SetTableView(JobPlanningLine);
    JobPlanningLinesPage.Run();
end;""")

h(doc, '6.2  Page 50662 "Workorder Card" - "Prepare Invoice Lines"', 2)
p(doc,
  'A single promoted action calling PrepareInvoiceLines(JobNo, JobTaskNo) - a broad sweep '
  'of every posted, unlinked Day Planning row for that Work Order''s Project No./Project '
  'Task No., intentionally NOT scoped to a selection (there is nothing to select on a card '
  'page). Refreshes the embedded "Project Planning Lines Part" via CurrPage.Update() after.')
add_code(doc,
r"""trigger OnAction()
var
    JobInvoicePrepMgt: Codeunit "Job Invoice Prep. Mgt.";
    LinesCreated, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount: Integer;
begin
    LinesCreated := JobInvoicePrepMgt.PrepareInvoiceLines(Rec."Project No.", Rec."Project Task No.", ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount);
    CurrPage.Update();
    Message(JobInvoicePrepMgt.FormatResultMessage(LinesCreated, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount));
end;""")

h(doc, '6.3  Page 50681 "Job Ledger Invoice Link" [NEW]', 2)
tbl = add_table(doc, [5, 12], 'Property', 'Value')
p681_rows = [
    ('PageType', 'List'),
    ('Editable / InsertAllowed / ModifyAllowed', 'false'),
    ('DeleteAllowed', 'false - see Section 5.2 for why manual deletion is blocked'),
    ('Fields shown', 'Job Ledger Entry No., Job No., Job Task No., Invoice Job Planning Line No. (drilldown), Skill Code'),
]
for i, r in enumerate(p681_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()
p(doc, 'The "Invoice Job Planning Line No." field has an OnDrillDown that opens the actual '
       'Job Planning Line it points to, via the standard "Job Planning Lines Part" filtered '
       'to that one Job No./Job Task No./Line No., with a graceful message if the target no '
       'longer exists.')
add_code(doc,
r"""trigger OnDrillDown()
var
    JobPlanningLine: Record "Job Planning Line";
    JobPlanningLinesPart: Page "Job Planning Lines Part";
begin
    JobPlanningLine.SetRange("Job No.", Rec."Job No.");
    JobPlanningLine.SetRange("Job Task No.", Rec."Job Task No.");
    JobPlanningLine.SetRange("Line No.", Rec."Invoice Job Planning Line No.");
    if JobPlanningLine.FindFirst() then begin
        JobPlanningLinesPart.SetRecord(JobPlanningLine);
        JobPlanningLinesPart.SetTableView(JobPlanningLine);
        JobPlanningLinesPart.Run();
    end else
        Message('The Project Planning Line that this usage entry was rolled into could not be found. It may have been deleted.');
end;""")

h(doc, '6.4  Page Extension 50626 "Skill Codes Opt." [NEW]', 2)
p(doc,
  'Extends the STANDARD base-app "Skill Codes" page (namespace Microsoft.Service.Setup) - '
  'the real master-data maintenance page for the Skill Code table - with the new "Invoice '
  'Resource No." field. A separate, pre-existing custom page in this app, "Opti Skill '
  'Codes" (page 50646), was deliberately NOT used for this: it is SourceTableTemporary = '
  'true, used only as a temp lookup/picker, so any field added there would never persist '
  'to real Skill Code records.')
add_code(doc,
r"""pageextension 50626 "Skill Codes Opt." extends "Skill Codes"
{
    layout
    {
        addafter(Description)
        {
            field("Invoice Resource No."; Rec."Invoice Resource No.")
            {
                ApplicationArea = All;
                ToolTip = 'Specifies the resource that should be used when invoicing usage recorded under this skill.';
            }
        }
    }
}""")

h(doc, '6.5  Planning Role Center - Navigation', 2)
tbl = add_table(doc, [4, 4, 9], 'Nav Group', 'New Action', 'Target')
nav_rows = [
    ('Resource', 'Skill Codes', 'Standard base-app "Skill Codes" page (same one extended in 6.4).'),
    ('Project Planning', 'Job Ledger Invoice Link', 'Page 50681 (Section 6.3), for direct traceability browsing/audit.'),
]
for i, r in enumerate(nav_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

# ── 7. REPORTS ─────────────────────────────────────────────────────────────────

h(doc, '7. Reports')

h(doc, '7.1  Report 50606 "Job Planning Line Usage Detail" [NEW]', 2)
p(doc,
  'The "customer detail hours report" from the design doc (Section 6). The invoice stays '
  'summarized per Skill, but this report un-collapses it: given one or more Job Planning '
  'Lines, it walks Job Ledger Invoice Link to find every posted usage entry rolled into '
  'each one, printing Posting Date / Resource No. / Description / Quantity / UOM per entry, '
  'plus a per-line summary grouped by Resource No. and total hours.')
tbl = add_table(doc, [5, 12], 'Property', 'Value')
r606_rows = [
    ('ProcessingOnly', 'false - this report displays data, it does not process anything'),
    ('UsageCategory / ApplicationArea', 'ReportsAndAnalysis / All - auto-listed under Tell Me / All Reports'),
    ('RequestFilterFields', 'Job No., Job Task No., Line No. (on the Job Planning Line dataitem)'),
    ('Layout', 'BC auto-generated layout - no RDLC/Word layout defined, deliberately kept simple per the Release 1 design'),
    ('Entry point', 'None yet - reachable today only via Tell Me / All Reports, no page action currently calls it'),
]
for i, r in enumerate(r606_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

info_box(doc,
    'Report 50606 is a real, working, compiled object - it is not dead code. It is simply '
    'not wired to any button or menu item anywhere in the app yet, which is why grepping '
    'the codebase for its object name turns up only the report''s own file. It still shows '
    'up to users today because of two page properties on the report itself: '
    'UsageCategory = ReportsAndAnalysis and ApplicationArea = All. Those two properties are '
    'what make Business Central automatically list a report under "Tell Me" search and '
    'under the "All Reports" page/Departments > Reports & Analysis listing, WITHOUT any '
    'explicit action or RunObject pointing at it anywhere else in the app - this is '
    'standard BC platform behaviour for any report, not something specific to this feature. '
    'So today, a user can only reach it by typing its name into Tell Me, or browsing All '
    'Reports and running it manually (which opens its request page, letting them filter by '
    'Job No./Job Task No./Line No. and see the detail for whichever Job Planning Line(s) '
    'they filter to). What is genuinely missing is a dedicated "Usage Detail" / "Print" '
    'button placed directly on a page a user would already be looking at when they want '
    'this - for example on the "Job Planning Lines Part" embedded on the Work Order card, '
    'or on the standard Job Planning Lines page itself, filtered to the currently selected '
    'line(s). That wiring was intentionally left undone in Release 1 (it was flagged and '
    'offered during implementation, but not requested) and remains open for a future pass - '
    'see Section 9.',
    label='Why This Report Has No Caller Yet', color='FFF2CC')

h(doc, '7.2  Report 50607 "Prepare Proj. Planning Lines" [NEW]', 2)
p(doc,
  'A proper batch-report entry point (request page, ProcessingOnly = true) for "Prepare '
  'Invoice Lines", complementing the inline page-action version on the Day Plannings list. '
  'Does the entire batch of work once in OnPreDataItem by handing the DataItem''s own '
  '(request-page-filtered) Day Planning recordset straight into '
  'PrepareInvoiceLinesForSelection, then CurrReport.Break()s to skip the report engine''s '
  'own row-by-row iteration (there is nothing to print - this is a pure batch process).')
add_code(doc,
r"""report 50607 "Prepare Proj. Planning Lines"
{
    Caption = 'Prepare Project Planning Lines for Invoicing';
    ProcessingOnly = true;
    UsageCategory = None;
    ApplicationArea = All;
    Permissions = tabledata "Day Planning" = r;

    dataset
    {
        dataitem(DayPlanning; "Day Planning")
        {
            RequestFilterFields = "Job No.", "Job Task No.", "Skill", "Assigned Resource No.";
            DataItemTableView = where(Posted = const(true));

            trigger OnPreDataItem()
            var
                JobInvoicePrepMgt: Codeunit "Job Invoice Prep. Mgt.";
                LinesCreated, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount: Integer;
            begin
                LinesCreated := JobInvoicePrepMgt.PrepareInvoiceLinesForSelection(DayPlanning, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount);
                Message(JobInvoicePrepMgt.FormatResultMessage(LinesCreated, ProcessedCount, AlreadyLinkedCount, NotPostedCount, SkippedOtherCount));
                CurrReport.Break();
            end;
        }
    }

    requestpage
    {
        Caption = 'Prepare Project Planning Lines for Invoicing';
        layout { area(content) { } }
    }
}""")

# ── 8. DATA REPAIR ────────────────────────────────────────────────────────────

h(doc, '8. Data Repair - Report 50600 "RepairData"')
p(doc,
  'A pre-existing internal repair/maintenance report gained a new action, '
  'DeleteGeneratedInvoicePlanningLines, for a one-off cleanup of every Job Planning Line '
  'ever generated by this feature (e.g. lines produced by an earlier, since-corrected '
  'revision of the logic during initial rollout), so they can be regenerated correctly.')

tbl = add_table(doc, [4, 13], 'Pass', 'Logic')
repair_rows = [
    ('1', 'Loop "Job Ledger Invoice Link" company-wide (unfiltered) and collect the DISTINCT '
          'set of (Job No., Job Task No., Line No.) keys into a temporary Job Planning Line '
          'record, deduping via Get-then-Insert.'),
    ('2', 'Loop the temporary key set, Get() the REAL Job Planning Line for each key, and '
          'Delete(true) it. Using Delete(true) (RunTrigger = true) is essential: it fires the '
          'OnAfterDeleteEvent subscriber (Section 5.2), which automatically cleans up the '
          'matching Job Ledger Invoice Link rows - no direct deletion from that table is done '
          'here.'),
]
for i, r in enumerate(repair_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

add_code(doc,
r"""local procedure DeleteGeneratedInvoicePlanningLines(): Integer
var
    JobLedgerInvoiceLink: Record "Job Ledger Invoice Link";
    TempJobPlanningLine: Record "Job Planning Line" temporary;
    JobPlanningLine: Record "Job Planning Line";
    n: Integer;
begin
    if JobLedgerInvoiceLink.FindSet() then
        repeat
            if not TempJobPlanningLine.Get(JobLedgerInvoiceLink."Job No.", JobLedgerInvoiceLink."Job Task No.", JobLedgerInvoiceLink."Invoice Job Planning Line No.") then begin
                TempJobPlanningLine.Init();
                TempJobPlanningLine."Job No." := JobLedgerInvoiceLink."Job No.";
                TempJobPlanningLine."Job Task No." := JobLedgerInvoiceLink."Job Task No.";
                TempJobPlanningLine."Line No." := JobLedgerInvoiceLink."Invoice Job Planning Line No.";
                TempJobPlanningLine.Insert();
            end;
        until JobLedgerInvoiceLink.Next() = 0;

    if TempJobPlanningLine.FindSet() then
        repeat
            if JobPlanningLine.Get(TempJobPlanningLine."Job No.", TempJobPlanningLine."Job Task No.", TempJobPlanningLine."Line No.") then begin
                JobPlanningLine.Delete(true);
                n += 1;
            end;
        until TempJobPlanningLine.Next() = 0;
    exit(n);
end;""")

info_box(doc,
    'This is a company-wide, unfiltered cleanup - it deletes every generated Job Planning '
    'Line, not just those for one Job. Intended for one-off use during initial rollout/'
    'correction, not as a routine operation.',
    label='Scope Warning', color='FFF2CC')

# ── 9. KNOWN LIMITATIONS & OPEN ITEMS ─────────────────────────────────────────

h(doc, '9. Known Limitations & Open Items (carried from the Release 1 design doc)')
tbl = add_table(doc, [5, 12], 'Item', 'Status / Note')
open_items = [
    ('Unit Price source',
     'Resolved via standard Job Planning Line field-validation (resource price list / job '
     'price group / work type), not grouped separately - matches a manual entry''s own '
     'price resolution. Left open by the design doc; this is the Release 1 choice made.'),
    ('Skill source on Job Ledger Entry',
     'Resolved via the Day Planning row''s own Skill field, reached through "Opt. '
     'DayPlanning Line No." / "Job Entry No." traceability - not re-derived independently '
     'on the Job Ledger Entry.'),
    ('UOM handling', 'Included in the grouping key (Job No. + Job Task No. + Invoice Resource No. + UOM) since multiple UOMs are possible.'),
    ('Reversal/corrections', 'Not handled in Release 1 - negative usage entries and credit/correction flows are explicitly out of scope, per the design doc.'),
    ('Installments', 'Out of scope for Release 1; the design doc notes Release 2 can reuse the same link table and keys.'),
    ('"Project Planning Lines" navigation filter precision', 'See Section 6.1 info box - exact for a single Job/Job Task selection, a mild over-approximation for multi-pair selections.'),
    ('Report 50606 entry point', 'No page action currently launches it (Section 7.1) - reachable via Tell Me / All Reports only.'),
]
for i, r in enumerate(open_items):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

# ── 10. OBJECT ID QUICK REFERENCE ─────────────────────────────────────────────

h(doc, '10. Object ID Quick Reference')
tbl = add_table(doc, [3, 3, 11], 'Type', 'ID', 'Name')
quick_ref = [
    ('Table', '50614', 'Job Ledger Invoice Link'),
    ('Table Extension', '50609', 'Opt. Skill Code'),
    ('Codeunit', '50607', 'Job Invoice Prep. Mgt.'),
    ('Codeunit', '50603', 'EventSubs (extended)'),
    ('Report', '50606', 'Job Planning Line Usage Detail'),
    ('Report', '50607', 'Prepare Proj. Planning Lines'),
    ('Report', '50600', 'RepairData (extended)'),
    ('Page', '50630', 'Day Plannings (extended)'),
    ('Page', '50662', 'Workorder Card (extended)'),
    ('Page', '50681', 'Job Ledger Invoice Link'),
    ('Page Extension', '50626', 'Skill Codes Opt.'),
    ('Page', '50612', 'Planning Role Center (extended)'),
]
for i, r in enumerate(quick_ref):
    data_row(tbl, *r, shade=(i % 2 == 0))

# ── SAVE ───────────────────────────────────────────────────────────────────────

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Project_Planning_Lines_for_Invoicing.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
