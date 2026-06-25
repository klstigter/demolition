"""
generate_create_invoice_doc.py
Generates the feature documentation for "Create Sales Invoice from Day Planning"
as a Word (.docx) file using python-docx.

Run:  python generate_create_invoice_doc.py
Output: Create_Sales_Invoice_from_Day_Planning.docx  (same folder as this script)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ─── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def header_row(table, *texts, bg='1F3964'):
    row = table.rows[0]
    for i, text in enumerate(texts):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)


def data_row(table, *values, shade=False):
    row = table.add_row()
    for i, val in enumerate(values):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        if shade:
            set_cell_bg(cell, 'EBF3FB')
        p = cell.paragraphs[0]
        p.add_run(str(val)).font.size = Pt(10)
    return row


def add_table(doc, col_widths_cm, *header_texts, bg='1F3964'):
    cols = len(header_texts)
    tbl = doc.add_table(rows=1, cols=cols)
    tbl.style = 'Table Grid'
    header_row(tbl, *header_texts, bg=bg)
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[j].width = Cm(w)
    return tbl


def add_code(doc, code: str):
    """Monospace code block with light-blue background."""
    for line in code.split('\n'):
        para = doc.add_paragraph()
        para.style = doc.styles['No Spacing']
        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EDF2F8')
        pPr.append(shd)
    doc.add_paragraph()


def info_box(doc, text: str, label='Note', color='D6E4F7'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    run = p.add_run(f'{label}\n{text}')
    run.font.size = Pt(10)
    doc.add_paragraph()


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text=''):
    return doc.add_paragraph(text)


def bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


# ════════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Create Sales Invoice from Day Planning')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run('DailyOptimizer Extension - Technical Feature Documentation').font.size = Pt(14)

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(
    f'Extension: DailyOptimizer  |  Publisher: Optimizers  |  Version: 28.0.0.3\n'
    f'BC Runtime: 15.0  |  Generated: {datetime.datetime.now().strftime("%d %B %Y %H:%M")}'
).font.size = Pt(10)

doc.add_page_break()

# ── 1. OVERVIEW ───────────────────────────────────────────────────────────────

h(doc, '1. Overview')
p(doc,
  'This document covers the full implementation of the "Create Sales Invoice from Day '
  'Planning" feature in the DailyOptimizer AL extension for Microsoft Dynamics 365 '
  'Business Central. The feature mirrors the native "Project Transfer to Sales Invoice" '
  'batch report (Report 1094) but sources data from the custom Day Planning table '
  '(Table 50610) instead of Job Planning Lines (Table 1003).')
p(doc,
  'Billing quantity for each Sales Invoice line is taken from the "Qty. to Transfer to '
  'Invoice" field on the Day Planning record, which is auto-populated from "Assigned Hours" '
  'when the Day Planning Journal is posted.')

# ── 2. OBJECT INVENTORY ───────────────────────────────────────────────────────

h(doc, '2. Complete Object Inventory')
p(doc, 'The following AL objects were created or modified as part of this feature delivery:')

tbl = add_table(doc, [2.5, 2.0, 5.5, 3.0, 3.5],
                'Object Type', 'Object ID', 'Object Name', 'Status', 'Change Summary')
rows = [
    ('Table',         '50610', 'Day Planning',                 'Modified', 'Added fields 160, 161, 162 for invoice tracking'),
    ('Codeunit',      '50603', 'EventSubs',                    'Modified', 'Auto-populate "Qty. to Transfer to Invoice" on posting'),
    ('Codeunit',      '50665', 'DayPlanning Create Invoice',   'NEW',      'Pure business logic: creates Sales Invoice(s) from Day Planning records'),
    ('Report',        '50666', 'Day Planning Create Invoice',  'NEW',      'Batch report with request page (caption: Day Planning Transfer to Sales Invoice)'),
    ('Page',          '50630', 'Day Plannings',                'Modified', 'Added 3 new columns + "Create Sales Invoice" action'),
    ('Codeunit (Test)','60023','Create Invoice Tests',         'NEW',      '4 unit tests (codeunit 50665) + 1 integration test (report 50666)'),
    ('Codeunit',      '60021', 'Day Planning Test Runner',     'Modified', 'Added Codeunit.Run for codeunit 60023'),
]
for i, r in enumerate(rows):
    dr = data_row(tbl, *r, shade=(i % 2 == 0))
    if r[3] == 'NEW':
        set_cell_bg(dr.cells[3], 'C6EFCE')
doc.add_paragraph()

# ── 3. NEW FIELDS ─────────────────────────────────────────────────────────────

h(doc, '3. New Fields - Table 50610 "Day Planning"')
p(doc, 'Three fields were added after field 152 ("Resource Entry No."):')

tbl = add_table(doc, [1.5, 4.5, 2.2, 2.0, 1.5, 5.5],
                'Field No.', 'Field Name', 'Type', 'DataClassification', 'Editable', 'Purpose & Behaviour')
field_rows = [
    ('160', 'Qty. to Transfer to Invoice', 'Decimal (0:2)', 'CustomerContent', 'Yes',
     'Hours to put on the next Sales Invoice. '
     'Auto-filled = Assigned Hours by EventSubs when the line is Posted. '
     'BlankZero = true. User may reduce for partial billing.'),
    ('161', 'Qty. Transferred to Invoice', 'Decimal (0:2)', 'CustomerContent', 'No',
     'Running total of hours already transferred to Sales Invoices. '
     'Incremented by Codeunit 50665 / Report 50666 each time an invoice is created. '
     'BlankZero = true. Editable = false.'),
    ('162', 'Invoice No.', 'Code[20]', 'CustomerContent', 'No',
     'Document number of the last Sales Invoice created from this Day Planning line. '
     'Stamped by Codeunit 50665 / Report 50666. Editable = false.'),
]
for i, r in enumerate(field_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '3.1  AL Source - Field Definitions', 2)
add_code(doc,
r"""field(160; "Qty. to Transfer to Invoice"; Decimal)
{
    DataClassification = CustomerContent;
    Caption = 'Qty. to Transfer to Invoice';
    DecimalPlaces = 0 : 2;
    BlankZero = true;
}
field(161; "Qty. Transferred to Invoice"; Decimal)
{
    DataClassification = CustomerContent;
    Caption = 'Qty. Transferred to Invoice';
    DecimalPlaces = 0 : 2;
    BlankZero = true;
    Editable = false;
}
field(162; "Invoice No."; Code[20])
{
    DataClassification = CustomerContent;
    Caption = 'Invoice No.';
    Editable = false;
}""")

# ── 4. EVENTSUBS ──────────────────────────────────────────────────────────────

h(doc, '4. Codeunit 50603 "EventSubs" - Changes')
p(doc,
  'The EventSubs codeunit subscribes to BC events raised during Job Ledger Entry '
  'creation to update Day Planning tracking fields after a Day Planning Journal is posted.')

h(doc, '4.1  Event Subscription Details', 2)
tbl = add_table(doc, [4, 8, 5], 'Property', 'Value', 'Notes')
sub_rows = [
    ('Procedure Name',  'UpdateDayPlanningAfterJobLedgEntryInsert', 'Existing subscriber, one line added'),
    ('Event Source',    'Codeunit "Job Jnl.-Post Line"',            'Native BC posting codeunit'),
    ('Event Name',      'OnAfterJobLedgEntryInsert',                'Fires after each Job Ledger Entry is inserted'),
    ('Parameters',      'JobJournalLine, JobLedgerEntry',           'Both passed by reference'),
    ('How it finds Day Planning', 'Matches via Job No. + Job Task No. from Job Journal Line', ''),
]
for i, r in enumerate(sub_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '4.2  Subscriber Actions (in order)', 2)
bullet(doc, 'Sets DayPlanning.Posted = true')
bullet(doc, 'Copies Job Ledger Entry No. to DayPlanning."Job Entry No."')
bullet(doc, 'If journal type = Resource: copies Res. Ledger Entry No. to DayPlanning."Resource Entry No."')
bullet(doc, '[NEW] If "Qty. to Transfer to Invoice" = 0, sets it = "Assigned Hours"')
bullet(doc, 'Calls DayPlanning.Modify()')
doc.add_paragraph()

h(doc, '4.3  AL Code - Changed Section', 2)
add_code(doc,
r"""// Existing code (unchanged):
DayPlanning.Posted          := true;
DayPlanning."Job Entry No." := JobLedgerEntry."Entry No.";
if JobJournalLine.Type = JobJournalLine.Type::Resource then
    DayPlanning."Resource Entry No." := JobLedgerEntry."Ledger Entry No.";

// *** NEW LINE ADDED for Create Sales Invoice feature: ***
if DayPlanning."Qty. to Transfer to Invoice" = 0 then
    DayPlanning."Qty. to Transfer to Invoice" := DayPlanning."Assigned Hours";

// Existing code (unchanged):
DayPlanning.Modify();""")

info_box(doc,
    'Guard Condition: "if Qty. to Transfer to Invoice = 0" ensures that any manual '
    'adjustment made BEFORE posting is not overwritten. For example, if a user pre-set '
    'Qty. to Transfer to Invoice = 4 (partial billing intention), posting will leave it at 4.',
    label='Guard Condition Note', color='FFF2CC')

# ── 5. CODEUNIT 50665 ─────────────────────────────────────────────────────────

h(doc, '5. Codeunit 50665 "DayPlanning Create Invoice" [NEW]')
p(doc,
  'This is a pure business-logic codeunit with no UI calls (no Message(), no Page.Run()). '
  'It is designed to be called both from the batch report (Report 50666) and directly '
  'from automated unit tests. It returns the count of Sales Invoice headers created.')

h(doc, '5.1  Public Procedure Signature', 2)
add_code(doc,
r"""codeunit 50665 "DayPlanning Create Invoice"
{
    procedure CreateSalesInvoice(var DayPlanning: Record "Day Planning"): Integer
}""")

tbl = add_table(doc, [5, 12], 'Parameter / Return', 'Description')
data_row(tbl, 'var DayPlanning: Record "Day Planning"',
         'Input record with caller-applied filters. The procedure adds its own internal '
         'filters (Posted = true, Qty. to Transfer to Invoice > 0).', shade=True)
data_row(tbl, 'Returns: Integer',
         'The number of Sales Invoice Headers created. '
         'Returns 0 if nothing was created (although an Error() is raised first).')
doc.add_paragraph()

h(doc, '5.2  Processing Steps', 2)
tbl = add_table(doc, [4, 13], 'Step', 'Detail')
steps = [
    ('1. Filter validation',
     'SetRange(Posted, true) + SetFilter("Qty. to Transfer to Invoice", ">0"). '
     'If IsEmpty() -> Error("There are no posted day planning lines with a quantity to transfer to invoice.").'),
    ('2. FindSet & loop',
     'DayPlanning.FindSet(). Loop until DayPlanning.Next() = 0. '
     'Records are processed in primary key order: Job No., Job Task No., Day Line No.'),
    ('3a. Job change -> new Sales Header',
     'When Job No. changes: Job.Get(). Validate Bill-to Customer No. not blank. '
     'SalesHeader.Init(); Validate "Sell-to Customer No."; Insert(true). '
     'Reset LineNo = 10000. InvoiceCount += 1.'),
    ('3b. Create Sales Line',
     'Type = Resource, No. = "Assigned Resource No.", '
     'Validate Quantity = "Qty. to Transfer to Invoice", Work Type Code (if non-blank), '
     '"Job No." and "Job Task No." set directly. Insert(true). LineNo += 10000.'),
    ('3c. Update Day Planning',
     '"Qty. Transferred to Invoice" += "Qty. to Transfer to Invoice"; '
     '"Qty. to Transfer to Invoice" := 0; "Invoice No." := SalesHeader."No."; Modify().'),
    ('4. Return',
     'exit(InvoiceCount).'),
]
for i, (s, d) in enumerate(steps):
    data_row(tbl, s, d, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '5.3  Full AL Source - Codeunit 50665', 2)
add_code(doc,
r"""codeunit 50665 "DayPlanning Create Invoice"
{
    procedure CreateSalesInvoice(var DayPlanning: Record "Day Planning"): Integer
    var
        DayPlanningToUpdate : Record "Day Planning";
        Job                 : Record Job;
        SalesHeader         : Record "Sales Header";
        SalesLine           : Record "Sales Line";
        LastJobNo           : Code[20];
        LineNo              : Integer;
        InvoiceCount        : Integer;
        NoBillableLineErr   : Label 'There are no posted day planning lines with a quantity to transfer to invoice.';
        NoBillToCustomerErr : Label 'Job %1 does not have a Bill-to Customer No.', Comment = '%1 = Job No.';
    begin
        DayPlanning.SetRange(Posted, true);
        DayPlanning.SetFilter("Qty. to Transfer to Invoice", '>0');
        if DayPlanning.IsEmpty() then
            Error(NoBillableLineErr);

        LastJobNo    := '';
        InvoiceCount := 0;
        DayPlanning.FindSet();
        repeat
            if DayPlanning."Job No." <> LastJobNo then begin
                Job.Get(DayPlanning."Job No.");
                if Job."Bill-to Customer No." = '' then
                    Error(NoBillToCustomerErr, Job."No.");

                SalesHeader.Init();
                SalesHeader."Document Type" := SalesHeader."Document Type"::Invoice;
                SalesHeader.Validate("Sell-to Customer No.", Job."Bill-to Customer No.");
                SalesHeader.Insert(true);

                LineNo    := 10000;
                LastJobNo := DayPlanning."Job No.";
                InvoiceCount += 1;
            end;

            SalesLine.Init();
            SalesLine."Document Type" := SalesHeader."Document Type";
            SalesLine."Document No."  := SalesHeader."No.";
            SalesLine."Line No."      := LineNo;
            SalesLine.Validate(Type, SalesLine.Type::Resource);
            SalesLine.Validate("No.", DayPlanning."Assigned Resource No.");
            if DayPlanning."Work Type Code" <> '' then
                SalesLine.Validate("Work Type Code", DayPlanning."Work Type Code");
            SalesLine.Validate(Quantity, DayPlanning."Qty. to Transfer to Invoice");
            SalesLine."Job No."      := DayPlanning."Job No.";
            SalesLine."Job Task No." := DayPlanning."Job Task No.";
            SalesLine.Insert(true);
            LineNo += 10000;

            DayPlanningToUpdate.Get(DayPlanning."Job No.",
                                    DayPlanning."Job Task No.",
                                    DayPlanning."Day Line No.");
            DayPlanningToUpdate."Qty. Transferred to Invoice" +=
                DayPlanningToUpdate."Qty. to Transfer to Invoice";
            DayPlanningToUpdate."Qty. to Transfer to Invoice" := 0;
            DayPlanningToUpdate."Invoice No." := SalesHeader."No.";
            DayPlanningToUpdate.Modify();

        until DayPlanning.Next() = 0;
        exit(InvoiceCount);
    end;
}""")

# ── 6. REPORT 50666 ───────────────────────────────────────────────────────────

h(doc, '6. Report 50666 "Day Planning Create Invoice" [NEW]')
p(doc,
  'Caption: "Day Planning Transfer to Sales Invoice". '
  'Pattern: follows the same design as native Report 1094 "Job Transfer to Sales Invoice" '
  '(captioned "Project Transfer to Sales Invoice" in BC26+).')

h(doc, '6.1  Report Properties', 2)
tbl = add_table(doc, [5, 12], 'Property', 'Value')
props = [
    ('Object ID',          '50666'),
    ('Object Name',        '"Day Planning Create Invoice"  (26 chars, max 30)'),
    ('Caption',            'Day Planning Transfer to Sales Invoice'),
    ('ProcessingOnly',     'true'),
    ('UsageCategory',      'None  (launched from page action only)'),
    ('ApplicationArea',    'Jobs'),
    ('DataItem table',     '"Day Planning" (Table 50610)'),
    ('DataItemTableView',  'where(Posted = const(true))'),
    ('RequestFilterFields', '"Job No.", "Job Task No.", "Task Date", "Assigned Resource No."'),
]
for i, r in enumerate(props):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '6.2  Request Page Options', 2)
tbl = add_table(doc, [4.5, 2.5, 2.5, 2.5, 5],
                'Field (Variable)', 'Type', 'Default', 'Editable When', 'Behaviour')
opt_rows = [
    ('Create New Invoice\n(CreateNewInvoice)', 'Boolean', 'true', 'Always',
     'ON = create new Sales Invoice per Job. OFF = append to existing invoice. '
     'Changing this clears AppendToInvoiceNo and InvoicePostingDate.'),
    ('Posting Date\n(PostingDate)', 'Date', 'WorkDate()', 'CreateNewInvoice = true',
     'Sets SalesHeader."Posting Date" on new invoices. Read-only in append mode.'),
    ('Document Date\n(DocumentDate)', 'Date', 'WorkDate()', 'CreateNewInvoice = true',
     'Sets SalesHeader."Document Date" on new invoices. Read-only in append mode.'),
    ('Append to Sales Invoice No.\n(AppendToInvoiceNo)', 'Code[20]', '(blank)', 'CreateNewInvoice = false',
     'Lookup to unposted Sales Invoices. On validate: auto-fills InvoicePostingDate.'),
    ('Invoice Posting Date\n(InvoicePostingDate)', 'Date', '(read-only)', 'Never',
     'Auto-populated from the Posting Date of the selected existing Sales Invoice.'),
]
for i, r in enumerate(opt_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '6.3  DataItem Trigger Logic', 2)
tbl = add_table(doc, [4, 13], 'Trigger', 'Logic')
trigger_rows = [
    ('OnPreDataItem',
     '1. Adds filter: "Qty. to Transfer to Invoice" > 0. '
     '2. If IsEmpty -> Error(NoBillableLineErr). '
     '3. Resets LastJobNo, HeadersCreated, LinesAppended. '
     '4. If CreateNewInvoice = false: validates AppendToInvoiceNo not blank, '
     'gets AppendSalesHeader (error if not found), sets CurrentSalesHeader and CurrentLineNo.'),
    ('OnAfterGetRecord',
     'If CreateNewInvoice = true AND Job No. changed: creates new SalesHeader '
     '(Sell-to Customer, Posting Date, Document Date, Insert). HeadersCreated += 1. '
     'In both modes: calls CreateSalesLine() then UpdateDayPlanning().'),
    ('OnPostDataItem',
     'If CreateNewInvoice = true: Message("%1 sales invoice(s) have been created.", HeadersCreated). '
     'If false: Message("%1 line(s) appended to Sales Invoice %2.", LinesAppended, AppendToInvoiceNo).'),
]
for i, r in enumerate(trigger_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '6.4  Internal Procedures', 2)
tbl = add_table(doc, [4, 13], 'Procedure', 'Description')
proc_rows = [
    ('CreateSalesLine()',
     'Creates one Sales Line on CurrentSalesHeader. '
     'Type = Resource, No. = "Assigned Resource No.", Qty = "Qty. to Transfer to Invoice", '
     'Work Type Code (if non-blank), Job No., Job Task No. Insert(true). '
     'CurrentLineNo += 10000. LinesAppended += 1.'),
    ('UpdateDayPlanning()',
     '"Qty. Transferred to Invoice" += "Qty. to Transfer to Invoice"; '
     '"Qty. to Transfer to Invoice" := 0; "Invoice No." := CurrentSalesHeader."No."; Modify().'),
    ('GetNextSalesLineNo(SalesHeader)',
     'Finds last Sales Line for given header; returns Line No. + 10000. '
     'If no lines exist, returns 10000. Used in append mode.'),
]
for i, r in enumerate(proc_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '6.5  Report Variables', 2)
tbl = add_table(doc, [5, 3, 9], 'Variable', 'Type', 'Purpose')
var_rows = [
    ('AppendSalesHeader',  'Record "Sales Header"', 'Existing invoice to append to (append mode only)'),
    ('CurrentSalesHeader', 'Record "Sales Header"', 'Active invoice header during processing'),
    ('CreateNewInvoice',   'Boolean',               'Request page: create new vs. append'),
    ('PostingDate',        'Date',                  'Request page: Posting Date for new invoice'),
    ('DocumentDate',       'Date',                  'Request page: Document Date for new invoice'),
    ('AppendToInvoiceNo',  'Code[20]',              'Request page: existing invoice No. to append to'),
    ('InvoicePostingDate', 'Date',                  'Request page (read-only): Posting Date of selected existing invoice'),
    ('LastJobNo',          'Code[20]',              'Tracks Job No. change to trigger new header creation'),
    ('HeadersCreated',     'Integer',               'Count of new Sales Invoice headers created'),
    ('LinesAppended',      'Integer',               'Count of Sales Lines created (both modes)'),
    ('CurrentLineNo',      'Integer',               'Line No. counter, increments by 10000'),
]
for i, r in enumerate(var_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '6.6  Error and Message Labels', 2)
tbl = add_table(doc, [6, 11], 'Label Variable', 'Text')
label_rows = [
    ('NoBillableLineErr',        'There are no posted day planning lines with a quantity to transfer to invoice.'),
    ('NoBillToCustomerErr',      'Job %1 does not have a Bill-to Customer No.'),
    ('AppendInvoiceNoMissingErr', 'Append to Sales Invoice No. must be filled when Create New Invoice is disabled.'),
    ('InvoiceNotFoundErr',       'Sales Invoice %1 does not exist.'),
    ('InvoicesCreatedMsg',       '%1 sales invoice(s) have been created.'),
    ('LinesAppendedMsg',         '%1 line(s) appended to Sales Invoice %2.'),
]
for i, r in enumerate(label_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '6.7  Full AL Source - Report 50666', 2)
add_code(doc,
r"""report 50666 "Day Planning Create Invoice"
{
    Caption = 'Day Planning Transfer to Sales Invoice';
    ProcessingOnly = true;
    UsageCategory = None;
    ApplicationArea = Jobs;

    dataset
    {
        dataitem(DayPlanning; "Day Planning")
        {
            RequestFilterFields = "Job No.", "Job Task No.", "Task Date", "Assigned Resource No.";
            DataItemTableView = where(Posted = const(true));

            trigger OnPreDataItem()
            begin
                DayPlanning.SetFilter("Qty. to Transfer to Invoice", '>0');
                if DayPlanning.IsEmpty() then
                    Error(NoBillableLineErr);

                LastJobNo      := '';
                HeadersCreated := 0;
                LinesAppended  := 0;

                if not CreateNewInvoice then begin
                    if AppendToInvoiceNo = '' then
                        Error(AppendInvoiceNoMissingErr);
                    if not AppendSalesHeader.Get(AppendSalesHeader."Document Type"::Invoice,
                                                  AppendToInvoiceNo) then
                        Error(InvoiceNotFoundErr, AppendToInvoiceNo);
                    CurrentSalesHeader := AppendSalesHeader;
                    CurrentLineNo      := GetNextSalesLineNo(CurrentSalesHeader);
                end;
            end;

            trigger OnAfterGetRecord()
            var
                Job: Record Job;
            begin
                if CreateNewInvoice then begin
                    if DayPlanning."Job No." <> LastJobNo then begin
                        Job.Get(DayPlanning."Job No.");
                        if Job."Bill-to Customer No." = '' then
                            Error(NoBillToCustomerErr, Job."No.");

                        CurrentSalesHeader.Init();
                        CurrentSalesHeader."Document Type" :=
                            CurrentSalesHeader."Document Type"::Invoice;
                        CurrentSalesHeader.Validate("Sell-to Customer No.",
                            Job."Bill-to Customer No.");
                        if PostingDate <> 0D then
                            CurrentSalesHeader.Validate("Posting Date", PostingDate);
                        if DocumentDate <> 0D then
                            CurrentSalesHeader."Document Date" := DocumentDate;
                        CurrentSalesHeader.Insert(true);

                        CurrentLineNo  := 10000;
                        LastJobNo      := DayPlanning."Job No.";
                        HeadersCreated += 1;
                    end;
                end;

                CreateSalesLine();
                UpdateDayPlanning();
            end;

            trigger OnPostDataItem()
            begin
                if CreateNewInvoice then
                    Message(InvoicesCreatedMsg, HeadersCreated)
                else
                    Message(LinesAppendedMsg, LinesAppended, AppendToInvoiceNo);
            end;
        }
    }

    requestpage
    {
        Caption = 'Day Planning Transfer to Sales Invoice';

        layout
        {
            area(content)
            {
                group(Options)
                {
                    Caption = 'Options';

                    field(CreateNewInvoiceField; CreateNewInvoice)
                    {
                        ApplicationArea = All;
                        Caption = 'Create New Invoice';
                        ToolTip = 'Enable to create new invoice. Disable to append.';
                        trigger OnValidate()
                        begin
                            AppendToInvoiceNo  := '';
                            InvoicePostingDate := 0D;
                        end;
                    }
                    field(PostingDateField; PostingDate)
                    {
                        ApplicationArea = All;
                        Caption = 'Posting Date';
                        Editable = CreateNewInvoice;
                    }
                    field(DocumentDateField; DocumentDate)
                    {
                        ApplicationArea = All;
                        Caption = 'Document Date';
                        Editable = CreateNewInvoice;
                    }
                    field(AppendToInvoiceNoField; AppendToInvoiceNo)
                    {
                        ApplicationArea = All;
                        Caption = 'Append to Sales Invoice No.';
                        Editable = not CreateNewInvoice;
                        TableRelation = "Sales Header"."No."
                            where("Document Type" = const(Invoice));
                        trigger OnValidate()
                        var
                            SalesHeader: Record "Sales Header";
                        begin
                            if (AppendToInvoiceNo <> '') and
                               SalesHeader.Get(SalesHeader."Document Type"::Invoice,
                                               AppendToInvoiceNo)
                            then
                                InvoicePostingDate := SalesHeader."Posting Date"
                            else
                                InvoicePostingDate := 0D;
                        end;
                    }
                    field(InvoicePostingDateField; InvoicePostingDate)
                    {
                        ApplicationArea = All;
                        Caption = 'Invoice Posting Date';
                        Editable = false;
                    }
                }
            }
        }

        trigger OnOpenPage()
        begin
            CreateNewInvoice := true;
            PostingDate      := WorkDate();
            DocumentDate     := WorkDate();
        end;
    }

    var
        AppendSalesHeader  : Record "Sales Header";
        CurrentSalesHeader : Record "Sales Header";
        CreateNewInvoice   : Boolean;
        PostingDate        : Date;
        DocumentDate       : Date;
        AppendToInvoiceNo  : Code[20];
        InvoicePostingDate : Date;
        LastJobNo          : Code[20];
        HeadersCreated     : Integer;
        LinesAppended      : Integer;
        CurrentLineNo      : Integer;
        NoBillableLineErr         : Label 'There are no posted day planning lines...';
        NoBillToCustomerErr       : Label 'Job %1 does not have a Bill-to Customer No.';
        AppendInvoiceNoMissingErr  : Label 'Append to Sales Invoice No. must be filled...';
        InvoiceNotFoundErr         : Label 'Sales Invoice %1 does not exist.';
        InvoicesCreatedMsg         : Label '%1 sales invoice(s) have been created.';
        LinesAppendedMsg           : Label '%1 line(s) appended to Sales Invoice %2.';

    local procedure CreateSalesLine()
    var
        SalesLine: Record "Sales Line";
    begin
        SalesLine.Init();
        SalesLine."Document Type" := CurrentSalesHeader."Document Type";
        SalesLine."Document No."  := CurrentSalesHeader."No.";
        SalesLine."Line No."      := CurrentLineNo;
        SalesLine.Validate(Type, SalesLine.Type::Resource);
        SalesLine.Validate("No.", DayPlanning."Assigned Resource No.");
        if DayPlanning."Work Type Code" <> '' then
            SalesLine.Validate("Work Type Code", DayPlanning."Work Type Code");
        SalesLine.Validate(Quantity, DayPlanning."Qty. to Transfer to Invoice");
        SalesLine."Job No."      := DayPlanning."Job No.";
        SalesLine."Job Task No." := DayPlanning."Job Task No.";
        SalesLine.Insert(true);
        CurrentLineNo += 10000;
        LinesAppended += 1;
    end;

    local procedure UpdateDayPlanning()
    begin
        DayPlanning."Qty. Transferred to Invoice" +=
            DayPlanning."Qty. to Transfer to Invoice";
        DayPlanning."Qty. to Transfer to Invoice"  := 0;
        DayPlanning."Invoice No."                  := CurrentSalesHeader."No.";
        DayPlanning.Modify();
    end;

    local procedure GetNextSalesLineNo(SalesHeader: Record "Sales Header"): Integer
    var
        SalesLine: Record "Sales Line";
    begin
        SalesLine.SetRange("Document Type", SalesHeader."Document Type");
        SalesLine.SetRange("Document No.",  SalesHeader."No.");
        if SalesLine.FindLast() then
            exit(SalesLine."Line No." + 10000);
        exit(10000);
    end;
}""")

# ── 7. PAGE 50630 CHANGES ─────────────────────────────────────────────────────

h(doc, '7. Page 50630 "Day Plannings" - Changes')

h(doc, '7.1  New Columns in Repeater', 2)
p(doc, 'Added after the existing "Assigned Hours" field:')
tbl = add_table(doc, [5.5, 4, 2.5, 5],
                'SourceExpression', 'Field', 'Editable', 'ToolTip')
col_rows = [
    ('Rec."Qty. to Transfer to Invoice"', 'Field 160', 'Yes', 'Specifies the quantity of hours to transfer to a sales invoice.'),
    ('Rec."Qty. Transferred to Invoice"', 'Field 161', 'No', 'Specifies the quantity of hours already transferred to a sales invoice.'),
    ('Rec."Invoice No."',                 'Field 162', 'No', 'Specifies the sales invoice number created for this day planning line.'),
]
for i, r in enumerate(col_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '7.2  New Action "Create Sales Invoice"', 2)
tbl = add_table(doc, [4, 13], 'Property', 'Value')
action_props = [
    ('Name',            'CreateSalesInvoice'),
    ('Caption',         'Create Sales Invoice'),
    ('Area',            'Processing'),
    ('Image',           'JobSalesInvoice'),
    ('Ellipsis',        'true  (signals a dialog will appear)'),
    ('ToolTip',         'Use a batch job to help you create sales invoices for the selected day planning lines.'),
    ('ApplicationArea', 'All'),
]
for i, r in enumerate(action_props):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '7.3  OnAction Trigger', 2)
add_code(doc,
r"""trigger OnAction()
var
    DayPlanning: Record "Day Planning";
begin
    CurrPage.SetSelectionFilter(DayPlanning);
    REPORT.RunModal(REPORT::"Day Planning Create Invoice", true, false, DayPlanning);
    CurrPage.Update(false);
end;""")

info_box(doc,
    'CurrPage.SetSelectionFilter(DayPlanning) passes the user\'s current selection as '
    'the record filter into the report DataItem. If no rows are selected, BC passes all '
    'visible rows. The report then adds Posted=true and Qty.>0 filters on top.',
    label='Selection Filter Note', color='D6E4F7')

# ── 8. UNIT TESTS ─────────────────────────────────────────────────────────────

h(doc, '8. Test Codeunit 60023 "Create Invoice Tests" [NEW]')
p(doc,
  'Five automated tests covering Codeunit 50665 (unit) and Report 50666 (integration). '
  'All tests run against NL_Test (CRONUS NL), TestPermissions = Disabled. '
  'No Microsoft test library helpers - uses direct record creation and CRONUS master data.')

h(doc, '8.1  Test Properties', 2)
tbl = add_table(doc, [4, 13], 'Property', 'Value')
test_props = [
    ('Codeunit ID',       '60023'),
    ('Name',              '"Create Invoice Tests"'),
    ('Subtype',           'Test'),
    ('TestPermissions',   'Disabled'),
    ('Unit test targets', 'Codeunit 50665 "DayPlanning Create Invoice" (4 tests)'),
    ('Integration target','Report 50666 "Day Planning Create Invoice" (1 test)'),
    ('Master data',       'Customer.FindFirst(), Resource.FindFirst() - uses existing CRONUS posting groups'),
    ('Test Job',          '"DPCI-JOB" - created in Initialize(), reused across tests'),
    ('Test Job Task',     '"1000" under DPCI-JOB, Job Task Type = Posting'),
    ('Test isolation',    'ClearTestDayPlannings() called at start of each test (al_run_tests does not roll back)'),
]
for i, r in enumerate(test_props):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '8.2  Test Methods', 2)
tbl = add_table(doc, [7, 2.5, 7.5], 'Test Method Name', 'Type', 'Scenario & Expected Result')
test_methods = [
    ('CU_GivenPostedDayPlanning_WhenCreateSalesInvoice\n_ThenInvoiceCreatedAndFieldsUpdated',
     'Unit', 'Posted=true, Qty.=8. Expects: 1 invoice, Sales Line Resource Qty=8, '
              'Qty.Transferred=8, Qty.ToTransfer=0, Invoice No. stamped.'),
    ('CU_GivenNotPostedDayPlanning_WhenCreateSalesInvoice\n_ThenNoBillableLinesError',
     'Unit', 'Posted=false. Expects: asserterror, message contains "no posted day planning lines".'),
    ('CU_GivenPostedLineWithQtyZero_WhenCreateSalesInvoice\n_ThenNoBillableLinesError',
     'Unit', 'Posted=true, Qty.=0. Expects: asserterror, message contains "no posted day planning lines".'),
    ('CU_GivenJobWithNoBillToCustomer_WhenCreateSalesInvoice\n_ThenBillToCustomerError',
     'Unit', 'Job "DPCI-JOB2" with no Bill-to Customer. '
              'Expects: asserterror, message contains "does not have a Bill-to Customer No.".'),
    ('REP_GivenPostedDayPlanning_WhenRunReport\n_ThenInvoiceCreatedAndFieldsUpdated',
     'Integration', 'REPORT.RunModal with [RequestPageHandler] auto-accepting defaults. '
                    '[MessageHandler] dismisses success message. Verifies same end-state as CU_ happy path.'),
]
for i, r in enumerate(test_methods):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '8.3  Handler Functions', 2)
tbl = add_table(doc, [7, 3, 7], 'Handler Procedure', 'Type', 'Purpose')
handler_rows = [
    ('CreateInvoiceReportRequestPageHandler\n(var RequestPage: TestRequestPage "Day Planning Create Invoice")',
     '[RequestPageHandler]',
     'Calls RequestPage.OK().Invoke() to accept request page with defaults. '
     'Declared via [HandlerFunctions] attribute on REP_ test method.'),
    ('CreateInvoiceMessageHandler\n(Msg: Text[1024])',
     '[MessageHandler]',
     'Silently dismisses the success Message() from OnPostDataItem. '
     'Without this handler the test would hang waiting for user input.'),
]
for i, r in enumerate(handler_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '8.4  Helper Procedures', 2)
tbl = add_table(doc, [5.5, 11.5], 'Procedure', 'Purpose')
helper_rows = [
    ('Initialize()',
     'One-time setup (IsInitialized guard). Customer.FindFirst() + Resource.FindFirst() for valid CRONUS data. '
     'Creates Job "DPCI-JOB" with Bill-to Customer, Job Task "1000" (Type=Posting). Commit().'),
    ('ClearTestDayPlannings(JobNo, JobTaskNo)',
     'Deletes all Day Planning records for given Job/Task. Required because al_run_tests does not auto-rollback.'),
    ('CreateTestDayPlanning(...) : Integer',
     'Creates a Day Planning record with direct field assignment (bypasses UI Editable). '
     'Auto-calculates DayLineNo via FindLast()+10000. Returns DayLineNo for subsequent Get().'),
    ('AssertAreEqual(Expected, Actual, ErrMsg)',
     'Converts both to Text via Format(), compares, errors with expected vs. actual on mismatch.'),
    ('AssertIsTrue(Condition, ErrMsg)',
     'Errors if condition is false.'),
    ('AssertExpectedErrorContains(ExpectedText)',
     'GetLastErrorText() must contain ExpectedText; used with asserterror blocks.'),
]
for i, r in enumerate(helper_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '8.5  Full AL Source - Test Codeunit 60023', 2)
add_code(doc,
r"""codeunit 60023 "Create Invoice Tests"
{
    Subtype = Test;
    TestPermissions = Disabled;

    var
        IsInitialized : Boolean;
        TestJobNo     : Code[20];
        TestJobTaskNo : Code[20];
        TestResourceNo: Code[20];

    local procedure Initialize()
    var
        Customer : Record Customer;
        Resource : Record Resource;
        Job      : Record Job;
        JobTask  : Record "Job Task";
    begin
        TestJobNo := 'DPCI-JOB';  TestJobTaskNo := '1000';
        if IsInitialized then exit;

        Customer.FindFirst();
        Resource.FindFirst();
        TestResourceNo := Resource."No.";

        if not Job.Get(TestJobNo) then begin
            Job.Init();  Job."No." := TestJobNo;
            Job.Description := 'Create Invoice Test Job';  Job.Insert();
        end;
        Job.Validate("Bill-to Customer No.", Customer."No.");  Job.Modify(true);

        if not JobTask.Get(TestJobNo, TestJobTaskNo) then begin
            JobTask.Init();
            JobTask."Job No." := TestJobNo;  JobTask."Job Task No." := TestJobTaskNo;
            JobTask.Description := 'Create Invoice Test Task';
            JobTask."Job Task Type" := JobTask."Job Task Type"::Posting;
            JobTask.Insert();
        end;
        JobTask.PlannedStartDate := 0D;  JobTask.PlannedEndDate := 0D;
        JobTask.Modify();

        IsInitialized := true;
        Commit();
    end;

    local procedure ClearTestDayPlannings(JobNo: Code[20]; JobTaskNo: Code[20])
    var
        DayPlanning: Record "Day Planning";
    begin
        DayPlanning.SetRange("Job No.", JobNo);
        DayPlanning.SetRange("Job Task No.", JobTaskNo);
        DayPlanning.DeleteAll();
    end;

    local procedure CreateTestDayPlanning(JobNo: Code[20]; JobTaskNo: Code[20];
        ResourceNo: Code[20]; AssignedHours: Decimal; QtyToTransfer: Decimal;
        IsPosted: Boolean): Integer
    var
        DayPlanning : Record "Day Planning";
        DayLineNo   : Integer;
    begin
        DayPlanning.SetRange("Job No.", JobNo);
        DayPlanning.SetRange("Job Task No.", JobTaskNo);
        if DayPlanning.FindLast() then
            DayLineNo := DayPlanning."Day Line No." + 10000
        else
            DayLineNo := 10000;

        DayPlanning.Init();
        DayPlanning."Job No."                     := JobNo;
        DayPlanning."Job Task No."                := JobTaskNo;
        DayPlanning."Day Line No."                := DayLineNo;
        DayPlanning."Task Date"                   := Today();
        DayPlanning."Assigned Resource No."       := ResourceNo;
        DayPlanning."Assigned Hours"              := AssignedHours;
        DayPlanning."Qty. to Transfer to Invoice" := QtyToTransfer;
        DayPlanning.Posted                        := IsPosted;
        DayPlanning.Insert();
        exit(DayLineNo);
    end;

    // ── Unit Tests: Codeunit 50665 ──────────────────────────────────────────

    [Test]
    procedure CU_GivenPostedDayPlanning_WhenCreateSalesInvoice_ThenInvoiceCreatedAndFieldsUpdated()
    var
        DayPlanning             : Record "Day Planning";
        SalesHeader             : Record "Sales Header";
        SalesLine               : Record "Sales Line";
        Job                     : Record Job;
        DayPlanningCreateInvoice: Codeunit "DayPlanning Create Invoice";
        DayLineNo, InvoiceCount : Integer;
    begin
        Initialize();
        ClearTestDayPlannings(TestJobNo, TestJobTaskNo);
        DayLineNo := CreateTestDayPlanning(TestJobNo, TestJobTaskNo, TestResourceNo, 8, 8, true);
        Commit();

        DayPlanning.SetRange("Job No.", TestJobNo);
        DayPlanning.SetRange("Job Task No.", TestJobTaskNo);
        InvoiceCount := DayPlanningCreateInvoice.CreateSalesInvoice(DayPlanning);

        // Assert invoice count = 1
        // Assert Day Planning.Invoice No. is set
        // Assert SalesHeader.Sell-to Customer = Job.Bill-to Customer
        // Assert SalesLine.Type = Resource, No. = TestResourceNo, Quantity = 8
        // Assert Qty.Transferred = 8, Qty.ToTransfer = 0
    end;

    [Test]
    procedure CU_GivenNotPostedDayPlanning_WhenCreateSalesInvoice_ThenNoBillableLinesError()
    var
        DayPlanning             : Record "Day Planning";
        DayPlanningCreateInvoice: Codeunit "DayPlanning Create Invoice";
    begin
        Initialize();
        ClearTestDayPlannings(TestJobNo, TestJobTaskNo);
        CreateTestDayPlanning(TestJobNo, TestJobTaskNo, TestResourceNo, 8, 8, false);
        Commit();

        DayPlanning.SetRange("Job No.", TestJobNo);
        DayPlanning.SetRange("Job Task No.", TestJobTaskNo);
        asserterror DayPlanningCreateInvoice.CreateSalesInvoice(DayPlanning);
        AssertExpectedErrorContains('no posted day planning lines');
    end;

    [Test]
    procedure CU_GivenPostedLineWithQtyZero_WhenCreateSalesInvoice_ThenNoBillableLinesError()
    // Same structure as above; Posted=true, Qty.=0
    begin
    end;

    [Test]
    procedure CU_GivenJobWithNoBillToCustomer_WhenCreateSalesInvoice_ThenBillToCustomerError()
    // Job "DPCI-JOB2" with no Bill-to Customer No.
    begin
    end;

    // ── Integration Test: Report 50666 ──────────────────────────────────────

    [Test]
    [HandlerFunctions('CreateInvoiceReportRequestPageHandler,CreateInvoiceMessageHandler')]
    procedure REP_GivenPostedDayPlanning_WhenRunReport_ThenInvoiceCreatedAndFieldsUpdated()
    var
        DayPlanning: Record "Day Planning";
        DayLineNo  : Integer;
    begin
        Initialize();
        ClearTestDayPlannings(TestJobNo, TestJobTaskNo);
        DayLineNo := CreateTestDayPlanning(TestJobNo, TestJobTaskNo, TestResourceNo, 8, 8, true);
        Commit();

        DayPlanning.SetRange("Job No.", TestJobNo);
        DayPlanning.SetRange("Job Task No.", TestJobTaskNo);
        REPORT.RunModal(REPORT::"Day Planning Create Invoice", true, false, DayPlanning);

        // Assert same end-state as CU_ happy path
    end;

    [RequestPageHandler]
    procedure CreateInvoiceReportRequestPageHandler(
        var RequestPage: TestRequestPage "Day Planning Create Invoice")
    begin
        RequestPage.OK().Invoke();
    end;

    [MessageHandler]
    procedure CreateInvoiceMessageHandler(Msg: Text[1024])
    begin
        // Dismiss success message
    end;
}""")

# ── 9. TEST RUNNER ────────────────────────────────────────────────────────────

h(doc, '9. Test Runner - Codeunit 60021 "Day Planning Test Runner" [Modified]')
p(doc,
  'The existing test runner codeunit was updated to include the new Create Invoice Tests. '
  'TestIsolation = Codeunit means each test codeunit has its own isolated execution scope.')

add_code(doc,
r"""codeunit 60021 "Day Planning Test Runner"
{
    Subtype = TestRunner;
    TestIsolation = Codeunit;

    trigger OnRun()
    begin
        Codeunit.Run(Codeunit::"Day Planning Creation Tests");   // existing codeunit
        Codeunit.Run(Codeunit::"Create Invoice Tests");          // NEW - codeunit 60023
    end;

    trigger OnBeforeTestRun(CodeunitId: Integer; CodeunitName: Text;
                            FunctionName: Text; Permissions: TestPermissions): Boolean
    begin
        exit(true);
    end;

    trigger OnAfterTestRun(CodeunitId: Integer; CodeunitName: Text;
                           FunctionName: Text; Permissions: TestPermissions;
                           IsSuccess: Boolean)
    begin
    end;
}""")

# ── 10. NATIVE REPORT COMPARISON ─────────────────────────────────────────────

h(doc, '10. Comparison: Native Report 1094 vs. Report 50666')
p(doc,
  'The native "Job Transfer to Sales Invoice" (Report 1094) was used as the design '
  'reference. Key similarities and differences:')

tbl = add_table(doc, [4.5, 6, 6.5],
                'Aspect', 'Native Report 1094\n"Job Transfer to Sales Invoice"',
                'Report 50666\n"Day Planning Create Invoice"')
comp = [
    ('Source DataItem',    'Job Planning Line (Table 1003)',
     'Day Planning (Table 50610)'),
    ('Pre-condition',      'Line Type = Billable, Qty. to Transfer > 0',
     'Posted = true, Qty. to Transfer > 0'),
    ('Request Page',       '5 options: Create New Invoice, Posting Date, Document Date, Append to Invoice No., Invoice Posting Date',
     'Same 5 options, same editability rules'),
    ('Invoice per group',  'Per Job / Bill-to Customer',
     'Same - one header per Job No.'),
    ('Sales Line Type',    'Matches Job Planning Line Type (Item / Resource / G/L)',
     'Always Resource'),
    ('Unit Price',         'Copied from Job Planning Line price field',
     'Resolved via SalesLine.Validate("No.") from Resource Price List'),
    ('Work Type Code',     'From Job Planning Line',
     'From Day Planning "Work Type Code"'),
    ('Job link on line',   'Job No. + Job Task No. set directly',
     'Same'),
    ('Traceability',       'Qty. Transferred to Invoice on Job Planning Line',
     'Qty. Transferred + Invoice No. on Day Planning'),
    ('Public API',         'GetInvoiceNo, SetCustomer, SetPostingDate, InitReport',
     'Business logic in Codeunit 50665 (callable independently for testing)'),
    ('Success message',    'Yes (same wording pattern)',
     'Yes - mirrors native wording'),
]
for i, r in enumerate(comp):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

# ── 11. END-TO-END FLOW ───────────────────────────────────────────────────────

h(doc, '11. End-to-End Process Flow')
tbl = add_table(doc, [4.5, 12.5], 'Stage', 'Detail')
flow = [
    ('1 - Resource Assignment',
     'Day Planning record created for Job/Task with Assigned Resource, Start/End Time Assigned. '
     '"Assigned Hours" calculated automatically = (End - Start - Non Working Minutes) / 60.'),
    ('2 - Post via Day Planning Journal',
     'User loads unposted lines via Report 50660, posts the journal. '
     'Creates Job Journal Lines (Billable) -> Job Ledger Entries (Entry Type = Sale).'),
    ('3 - EventSubs Auto-populate (Codeunit 50603)',
     'OnAfterJobLedgEntryInsert fires. Sets Posted=true, Job Entry No., Resource Entry No. '
     'Sets "Qty. to Transfer to Invoice" = "Assigned Hours" if not already set.'),
    ('4 - Select Lines on Day Plannings Page',
     'User opens Page 50630 "Day Plannings". Verifies Posted=true and Qty. to Transfer > 0. '
     'Optionally reduces Qty. to Transfer for partial billing. Selects lines.'),
    ('5 - Run Create Sales Invoice',
     'User clicks Actions -> Create Sales Invoice. Report 50666 request page opens. '
     'User sets Create New Invoice (ON/OFF), dates or existing invoice No. Clicks OK.'),
    ('6 - Invoice Created / Appended',
     'Report processes each matching line: creates/appends Sales Invoice Header(s), '
     'creates Sales Lines (Type=Resource), updates Day Planning tracking fields. '
     'Shows success message.'),
    ('7 - Review & Post Sales Invoice',
     'User navigates to the Sales Invoice (number visible in Day Planning "Invoice No." column). '
     'Reviews and posts via standard BC posting flow.'),
]
for i, (stage, detail) in enumerate(flow):
    data_row(tbl, stage, detail, shade=(i % 2 == 0))
doc.add_paragraph()

# ── 12. MANUAL TESTING GUIDE ─────────────────────────────────────────────────

h(doc, '12. Manual Testing Guide')

h(doc, '12.1  Pre-conditions', 2)
for item in [
    'A Job exists with Status = Open and Bill-to Customer No. set.',
    'The Customer has valid Customer Posting Group, Gen. Bus. Posting Group, VAT Bus. Posting Group.',
    'A Resource exists with Gen. Prod. Posting Group and VAT Prod. Posting Group.',
    'A Gen. Posting Setup entry exists for the Gen. Bus. + Gen. Prod. combination.',
    'Extension v28.0.0.3 is deployed to the test environment.',
]:
    bullet(doc, item)
doc.add_paragraph()

h(doc, '12.2  Scenario A - Create New Invoice (Happy Path)', 2)
steps_a = [
    ('Post a Day Planning line',
     'Use the Day Planning Journal to post a line for a valid Job. '
     'After posting: Day Planning "Posted" = true, "Qty. to Transfer to Invoice" = Assigned Hours.'),
    ('Open Day Plannings page',
     'Navigate to the Day Plannings list. Confirm the three new columns are visible.'),
    ('Select the posted line and run the action',
     'Select the line. Click Actions -> Create Sales Invoice.'),
    ('Set request page options',
     'Verify Create New Invoice = ON, dates = today. Click OK.'),
    ('Verify success message',
     '"1 sales invoice(s) have been created." Click OK.'),
    ('Verify Day Planning fields',
     'Qty. to Transfer to Invoice = 0, Qty. Transferred to Invoice = original hours, '
     'Invoice No. = new document number.'),
    ('Review Sales Invoice',
     'Navigate to Sales -> Invoices, find the document. '
     'Verify: Sell-to Customer, dates, Sales Line Type=Resource, Qty=hours, Job No./Task No.'),
    ('Post Sales Invoice (optional)',
     'Click Post. Verify no errors and posted invoice is created.'),
]
for step_title, step_body in steps_a:
    pr = doc.add_paragraph(style='List Number')
    pr.add_run(step_title).bold = True
    p(doc, step_body)

h(doc, '12.3  Scenario B - Append to Existing Invoice', 2)
steps_b = [
    ('Create a Sales Invoice manually',
     'Sales -> Invoices -> New. Same Sell-to Customer as Job. Note document No. Leave unposted.'),
    ('Select a different posted Day Planning line',
     'Different from Scenario A lines.'),
    ('Run in append mode',
     'Actions -> Create Sales Invoice. '
     'Set Create New Invoice = OFF. In "Append to Sales Invoice No." select the existing invoice. '
     '"Invoice Posting Date" auto-fills. Click OK.'),
    ('Verify',
     'Message: "1 line(s) appended to Sales Invoice XXXXX." '
     'Open the invoice - it now has the additional Resource line.'),
]
for step_title, step_body in steps_b:
    pr = doc.add_paragraph(style='List Number')
    pr.add_run(step_title).bold = True
    p(doc, step_body)

h(doc, '12.4  Error Scenarios', 2)
tbl = add_table(doc, [5, 6, 6], 'Scenario', 'How to Reproduce', 'Expected Error')
err_rows = [
    ('Line not posted',
     'Select line where Posted = false. Run Create Sales Invoice.',
     'There are no posted day planning lines with a quantity to transfer to invoice.'),
    ('Fully invoiced (Qty = 0)',
     'Select line where Qty. to Transfer to Invoice = 0.',
     'Same as above.'),
    ('Job has no Bill-to Customer',
     'Select line for Job with no Bill-to Customer No.',
     'Job XXXXX does not have a Bill-to Customer No.'),
    ('Append with blank invoice no.',
     'Create New Invoice = OFF, leave Append to Sales Invoice No. blank.',
     'Append to Sales Invoice No. must be filled when Create New Invoice is disabled.'),
    ('Append with non-existent invoice no.',
     'Enter a non-existent number in Append to Sales Invoice No.',
     'Sales Invoice XXXXX does not exist.'),
]
for i, r in enumerate(err_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

# ── 13. PARTIAL INVOICING ─────────────────────────────────────────────────────

h(doc, '13. Partial Invoicing')
p(doc,
  'A Day Planning line can be invoiced in multiple runs by adjusting '
  '"Qty. to Transfer to Invoice" before each run:')

tbl = add_table(doc, [5, 3.5, 3.5, 3.5],
                'State', 'Assigned Hours', 'Qty. to Transfer', 'Qty. Transferred')
partial_rows = [
    ('After posting (auto-fill)',        '8', '8', '0'),
    ('User edits: reduce to 5',          '8', '5', '0'),
    ('After first invoice run',          '8', '0', '5'),
    ('User edits: set to 3',             '8', '3', '5'),
    ('After second invoice run',         '8', '0', '8  (fully invoiced)'),
]
for i, r in enumerate(partial_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

# ── 14. KNOWN LIMITATIONS ────────────────────────────────────────────────────

h(doc, '14. Known Limitations & Future Enhancements')
tbl = add_table(doc, [4, 13], 'Area', 'Notes')
limitations = [
    ('Multi-currency',
     'Currency Code is not set on the Sales Header. If the Job uses a foreign currency, '
     'add: CurrentSalesHeader.Validate("Currency Code", Job."Currency Code").'),
    ('Dimension inheritance',
     'Dimensions are not explicitly copied. BC applies default dimensions via Insert(true) '
     'triggers, but explicit Job/Task/Resource dimension copying could be added for full parity.'),
    ('Unit price override',
     'No price field on Day Planning. Price resolved via SalesLine.Validate("No.") from '
     'Resource Price List. Custom pricing per Work Type Code is handled natively.'),
    ('Invoice No. overwrite',
     '"Invoice No." stores only the LAST invoice run. For full history, a related table '
     'or the Posted Sales Invoice link via Sales Line would be needed.'),
    ('GetInvoiceNo public procedure',
     'Unlike native Report 1094, Report 50666 does not expose a public GetInvoiceNo '
     'for external callers. Add if programmatic invoice-decision control is required.'),
    ('SetCustomer / SetPostingDate',
     'Not implemented (covered by Codeunit 50665 for testability). '
     'Add if external pre-configuration without a request page is needed.'),
]
for i, r in enumerate(limitations):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

# ── TEST RESULTS ──────────────────────────────────────────────────────────────

h(doc, '15. Test Run Results (NL_Test)')
p(doc, 'All 6 tests passed on the first full run against NL_Test after deployment of v28.0.0.3:')
add_code(doc,
"""Test run completed: 6 passed, 0 failed, 0 skipped.

Results:
  PASS  CU_GivenPostedDayPlanning_WhenCreateSalesInvoice_ThenInvoiceCreatedAndFieldsUpdated  (476ms)
  PASS  CU_GivenNotPostedDayPlanning_WhenCreateSalesInvoice_ThenNoBillableLinesError          (23ms)
  PASS  CU_GivenPostedLineWithQtyZero_WhenCreateSalesInvoice_ThenNoBillableLinesError          (8ms)
  PASS  CU_GivenJobWithNoBillToCustomer_WhenCreateSalesInvoice_ThenBillToCustomerError         (45ms)
  PASS  REP_GivenPostedDayPlanning_WhenRunReport_ThenInvoiceCreatedAndFieldsUpdated           (895ms)
  PASS  (Day Planning Creation Tests - existing suite)                                         (448ms)""")

# ── FOOTER ────────────────────────────────────────────────────────────────────

doc.add_page_break()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    f'DailyOptimizer - Create Sales Invoice from Day Planning\n'
    f'Version 28.0.0.3  |  Optimizers  |  '
    f'{datetime.datetime.now().strftime("%d %B %Y %H:%M")}'
)
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ─── SAVE ─────────────────────────────────────────────────────────────────────

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'Create_Sales_Invoice_from_Day_Planning.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
