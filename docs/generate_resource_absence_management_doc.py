"""
generate_resource_absence_management_doc.py
Generates the technical design document for "Resource Absence Management"
(review draft - no code implemented yet) as a Word (.docx) file using python-docx.

Run:  python generate_resource_absence_management_doc.py
Output: Resource Absence Management.docx  (same folder as this script)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ─── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def header_row(table, *texts, bg='1F3964'):
    row = table.rows[0]
    for i, text in enumerate(texts):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)


def data_row(table, *values, shade=False):
    row = table.add_row()
    for i, val in enumerate(values):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        if shade:
            set_cell_bg(cell, 'EBF3FB')
        p = cell.paragraphs[0]
        p.add_run(str(val)).font.size = Pt(10)
    return row


def add_table(doc, col_widths_cm, *header_texts, bg='1F3964'):
    cols = len(header_texts)
    tbl = doc.add_table(rows=1, cols=cols)
    tbl.style = 'Table Grid'
    header_row(tbl, *header_texts, bg=bg)
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[j].width = Cm(w)
    return tbl


def add_code(doc, code: str):
    """Monospace code block with light-blue background."""
    for line in code.split('\n'):
        para = doc.add_paragraph()
        para.style = doc.styles['No Spacing']
        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EDF2F8')
        pPr.append(shd)
    doc.add_paragraph()


def info_box(doc, text: str, label='Note', color='D6E4F7'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    run = p.add_run(f'{label}\n{text}')
    run.font.size = Pt(10)
    doc.add_paragraph()


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text=''):
    return doc.add_paragraph(text)


def bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


# ════════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Resource Absence Management')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run('DailyOptimizer Extension - Technical Design (Review Draft, no code implemented yet)').font.size = Pt(14)

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(
    f'Extension: DailyOptimizer  |  Publisher: Optimizers  |  Version: 28.0.0.3\n'
    f'BC Runtime: 15.0  |  Generated: {datetime.datetime.now().strftime("%d %B %Y %H:%M")}'
).font.size = Pt(10)

doc.add_paragraph()
info_box(doc,
    'This document is a REVISED REVIEW DRAFT. AL code exists for the PRIOR (Card-based) '
    'design iteration, which this revision supersedes with a journal-style worksheet + '
    'Post model (Section 2). Section 4''s Migration Note lists exactly what needs '
    'reworking, and Section 9 lists the open questions specific to this revision that '
    'should be confirmed before that rework starts.',
    label='Status', color='FFF2CC')

doc.add_page_break()

# ── 1. BACKGROUND & OBJECTIVE ─────────────────────────────────────────────────

h(doc, '1. Background & Objective')
p(doc,
  'Microsoft Dynamics 365 Business Central does not support Absence tracking for '
  'Resources. Absence functionality only exists in the Human Resources module, and is '
  'limited to Employees (table 5202 "Employee Absence", reason table 5206 "Cause of '
  'Absence"). Resources - the entity DailyOptimizer schedules capacity and Day Planning '
  'against - have no equivalent.')
p(doc,
  'Objective: develop a Resource Absence feature so that Resource capacity calculations '
  'account for absence hours, without duplicating or replacing the existing standard '
  'Resource Capacity infrastructure (table 160 "Res. Capacity Entry", the "Update '
  'Resource Capacity" wizard, and the capacity aggregation already used throughout this '
  'extension - see Section 8).')

# ── 2. DESIGN SUMMARY ─────────────────────────────────────────────────────────

h(doc, '2. Design Summary')
p(doc,
  'Two-stage design: absence entries are never written directly into table 160 "Res. '
  'Capacity Entry" by the user. Instead, a new UNPOSTED staging table, "Register '
  'Absence" (Section 3.5), holds a journal-style worksheet where a user can key in '
  'several absence lines - potentially for different resources and dates - before '
  'committing anything. A "Post" action then validates each line and, only at that '
  'point, creates the real ledger row in table 160. This mirrors the standard BC '
  'journal pattern (enter many lines, review, then Post) rather than a direct-entry '
  'Card, so a half-finished or miskeyed batch never touches the actual capacity ledger.')
p(doc,
  'The POSTED side of the design is unchanged from the previous iteration: table 160 '
  'is still extended (not replaced) so every existing SUM/FlowField/query that already '
  'reads "Res. Capacity Entry".Capacity keeps working with zero code changes - it nets '
  'Capacity against Absence automatically, because Absence is stored as a NEGATIVE value '
  'in that same Capacity field. Only the ENTRY MECHANISM changes; the ledger shape and '
  'all of Section 5''s validation rules carry over unchanged, just moved from "runs on '
  'Card insert" to "runs on Post".')

tbl = add_table(doc, [4, 13], 'Rule', 'Detail')
rules = [
    ('Unposted staging, then Post', 'New table "Register Absence" (Section 3.5) + worksheet page (Section 6.4) hold not-yet-committed lines. Nothing is written to table 160 until the user clicks Post.'),
    ('New "Type" field (posted ledger)', 'Enum "Res. Capacity Entry Type": 0 = Capacity, 1 = Absence. All existing table 160 rows default to 0 (Capacity) - fully backward compatible.'),
    ('Absence = negative Capacity (posted ledger)', 'A posted Absence entry reuses the standard "Start Time"/"End Time" fields and writes its Hours into the same "Capacity" field, but negated.'),
    ('Existing-capacity guard', 'A line cannot be posted for a Resource/Date that has no Capacity row already - checked both as a visible field on the worksheet and inside the posting codeunit - see Section 5.'),
]
for i, r in enumerate(rules):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '2.1  Worked Example (Posted State)', 2)
p(doc, 'After the worksheet line below has been posted, table 160 contains:')
tbl = add_table(doc, [1.5, 5, 3, 3], '#', 'Date', 'Type', 'Hours')
example = [
    ('1', '1 Aug 2026', 'Capacity', '5'),
    ('2', '1 Aug 2026', 'Capacity', '7'),
    ('3', '1 Aug 2026', 'Absence', '-4'),
]
for i, r in enumerate(example):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()
info_box(doc,
    'Total Resource Capacity for 1 Aug 2026 = 5 + 7 + (-4) = 8. No separate "net capacity" '
    'calculation is introduced - the existing SUM(Capacity) already produces this result.',
    label='Result', color='C6EFCE')

# ── 3. DATA MODEL CHANGES ─────────────────────────────────────────────────────

h(doc, '3. Data Model Changes')
p(doc,
  'All new fields are added to the existing table extension 50606 "ResCapacityEntry '
  'Opt" (which already extends table 160 "Res. Capacity Entry" with "Start Time", '
  '"End Time", "Duplicate Id" and the "Requested Hours" FlowField - see the extract in '
  'Section 3.2). No new base table is required for the capacity/absence ledger itself.')

h(doc, '3.1  New Fields on Table Extension 50606', 2)
tbl = add_table(doc, [1.3, 4.8, 4.0, 8.9],
                'Field No.', 'Field Name', 'Type', 'Purpose')
field_rows = [
    ('50605', 'Type', 'Enum "Res. Capacity Entry Type"',
     'Capacity (0, default) or Absence (1). Drives every validation rule in Section 5.'),
    ('50606', 'Absence Reason Code', 'Code[10]',
     'TableRelation to table 5206 "Cause of Absence" (the exact reason table standard '
     'Employee Absence already uses - reused as-is, no new setup table). Mandatory on '
     'Absence rows, blank on Capacity rows.'),
]
for i, r in enumerate(field_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

info_box(doc,
    'No separate "Capacity Entry No." link field is needed. Because Absence rows are '
    'ordinary rows in the same table 160, they are already identified by the standard '
    '"Resource No." + Date (+ "Duplicate Id") combination - the same combination the '
    'existing-capacity guard (Section 5.1) validates against and Query 50604 groups by. '
    'A stored back-reference to the specific Capacity row would be redundant data that '
    'could itself go stale (e.g. if that Capacity row is later deleted/regenerated).',
    label='Simplification', color='C6EFCE')

h(doc, '3.2  Existing Table Extension 50606 (for reference, unchanged fields shown)', 2)
p(doc, 'Current state of src/tableext/tableext_50606_ResourceCapacityEntry.al, which the fields above will be added to:')
add_code(doc,
r"""tableextension 50606 "ResCapacityEntry Opt" extends "Res. Capacity Entry"
{
    fields
    {
        field(50600; "Start Time"; Time) { DataClassification = ToBeClassified; }
        field(50601; "End Time"; Time) { DataClassification = ToBeClassified; }
        field(50602; "Duplicate Id"; Integer)
        {
            DataClassification = ToBeClassified;
            InitValue = 1;
        }
        field(50604; "Requested Hours"; Decimal)
        {
            Editable = false;
            fieldclass = FlowField;
            calcformula = sum("Day Planning"."Assigned Hours"
                where("Assigned Resource No." = field("Resource No."),
                      "Task Date" = field("Date")));
        }

        // --- new fields proposed by this design (50605-50606) ---
    }
}""")

h(doc, '3.3  New Enum 50682 "Res. Capacity Entry Type"', 2)
add_code(doc,
r"""enum 50682 "Res. Capacity Entry Type"
{
    Extensible = true;

    value(0; Capacity) { Caption = 'Capacity'; }
    value(1; Absence) { Caption = 'Absence'; }
}""")

h(doc, '3.4  Base Table 160 "Res. Capacity Entry" (unchanged, for reference)', 2)
p(doc, 'Confirmed via symbol search - the standard fields and keys this design builds on:')
tbl = add_table(doc, [4, 5, 8], 'Key', 'Fields', 'Relevance to this design')
base_keys = [
    ('Key1 (PK)', '"Entry No."', 'Auto-numbered Integer. Absence rows get their own Entry No. like any other row - no separate link field is stored (Section 3.1).'),
    ('Key2', '"Resource No.", Date', 'Used by the existing-capacity guard (Section 5) to find the matching Capacity row.'),
    ('Key3', '"Resource Group No.", Date', 'Unaffected by this design.'),
]
for i, r in enumerate(base_keys):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '3.5  New Table 50687 "Register Absence" [NEW] - Unposted Worksheet', 2)
p(doc,
  'The staging table behind the new worksheet page (Section 6.4). One row per not-yet-'
  'posted absence line. A row is DELETED once it posts successfully (standard journal '
  'behavior - a posted line does not linger in the worksheet) - table 160 is the only '
  'permanent record of a posted absence.')
tbl = add_table(doc, [1.3, 4.8, 4.0, 8.9],
                'Field No.', 'Field Name', 'Type', 'Purpose')
register_fields = [
    ('1', 'Line No.', 'Integer', 'Primary key. Plain incrementing line number, same convention as a journal line.'),
    ('2', 'Resource No.', 'Code[20]', 'TableRelation to Resource. Mandatory at Post (Section 5.8).'),
    ('3', 'Date', 'Date', 'The absence date. Mandatory at Post. (Captioned "Date" on the table; the worksheet page may caption the column "Absence Date" for clarity, matching the posted side''s field naming.)'),
    ('4', 'Absence Reason Code', 'Code[10]', 'TableRelation to table 5206 "Cause of Absence" - same reason table the posted ledger field uses. Mandatory at Post.'),
    ('5', 'Hours', 'Decimal', 'Positive hours the user is entering as absent. Mandatory (> 0) at Post.'),
    ('6', 'Existing Capacity', 'Decimal (FlowField)',
     'CalcFormula: sum("Res. Capacity Entry".Capacity where("Resource No." = field("Resource No."), '
     'Date = field(Date))) - deliberately NO Type filter, so this sums Capacity AND Absence rows '
     'together, giving the NET REMAINING capacity for the date (Capacity minus anything already '
     'posted as Absence), not just whether a Capacity row exists at all. Read-only, shown on the '
     'worksheet so the user can see BEFORE posting how much capacity is actually left. Mandatory '
     '(> 0) at Post - this is the worksheet-side surface of the existing-capacity guard (Section '
     '5.1), and also what correctly blocks posting once a resource''s capacity for that date has '
     'already been fully consumed by other posted absences, not only when it never existed.'),
]
for i, r in enumerate(register_fields):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

add_code(doc,
r"""table 50687 "Register Absence"
{
    Caption = 'Register Absence';
    DataClassification = CustomerContent;

    fields
    {
        field(1; "Line No."; Integer)
        {
            Caption = 'Line No.';
            DataClassification = SystemMetadata;
        }
        field(2; "Resource No."; Code[20])
        {
            Caption = 'Resource No.';
            DataClassification = CustomerContent;
            TableRelation = Resource;
        }
        field(3; Date; Date)
        {
            Caption = 'Date';
            DataClassification = CustomerContent;
        }
        field(4; "Absence Reason Code"; Code[10])
        {
            Caption = 'Absence Reason Code';
            DataClassification = CustomerContent;
            TableRelation = "Cause of Absence";
        }
        field(5; Hours; Decimal)
        {
            Caption = 'Hours';
            DataClassification = CustomerContent;
            DecimalPlaces = 0 : 2;
            BlankZero = true;
        }
        field(6; "Existing Capacity"; Decimal)
        {
            Caption = 'Existing Capacity';
            Editable = false;
            FieldClass = FlowField;
            // No Type filter: sums Capacity AND Absence together, so this reflects NET
            // remaining capacity, not just whether a Capacity row exists at all.
            CalcFormula = sum("Res. Capacity Entry".Capacity
                where("Resource No." = field("Resource No."),
                      Date = field(Date)));
            DecimalPlaces = 0 : 5;
            BlankZero = true;
        }
    }

    keys
    {
        key(PK; "Line No.")
        {
            Clustered = true;
        }
    }
}""")

# ── 4. OBJECT ID PLAN ─────────────────────────────────────────────────────────

h(doc, '4. New / Changed Object ID Plan')
p(doc,
  'The app''s id range is 50600-60700 (app.json). Highest IDs currently in use top out '
  'around 50686 in the low sub-range (the codeunit from the prior design iteration), so '
  'new objects below take the next free IDs starting at 50687.')

tbl = add_table(doc, [2.6, 1.6, 5.6, 2.2, 5.1],
                'Object Type', 'Object ID', 'Object Name', 'Status', 'Purpose')
rows = [
    ('Enum',            '50682', 'Res. Capacity Entry Type',   'NEW',      'Capacity / Absence classification on the posted ledger (Section 3.3)'),
    ('Table Extension', '50606', 'ResCapacityEntry Opt',       'Modified', '+ 2 fields: Type, Absence Reason Code (Section 3.1)'),
    ('Table',           '50687', 'Register Absence',           'NEW',      'Unposted worksheet staging table (Section 3.5)'),
    ('Page (Worksheet)', '50688', 'Register Absence',           'NEW',      'Journal-style entry page with the Post action (Section 6.4)'),
    ('Codeunit',        '50686', 'Resource Absence Mgt.',      'Modified', 'Validation logic unchanged; entry point changes from Card triggers to Post (Section 5)'),
    ('Page (List)',     '50684', 'Resource Absence List',      'Repurposed', 'Now READ-ONLY posted-history view only - no more New/Edit (Section 6.2)'),
    ('Page (Card)',     '50685', 'Resource Absence Card',      'REMOVED',  'Superseded by the worksheet - direct Card entry into table 160 is no longer part of the design (Section 6.3)'),
    ('Page Extension',  '50605', 'ResourceCard Opti',          'Modified', '"Absence" action retained, re-purposed to open the read-only history list (Section 6.1)'),
    ('Page',            '50629', 'Res. Capacity FactBox Part', 'Modified', '+ "Type" column so Capacity/Absence rows are distinguishable (Section 8.2)'),
    ('Page',            '50627', 'Resource Capacity Settings Opt', 'Modified', '"Update Capacity" wizard - CalcSums(Capacity) must add a Type = Capacity filter (Section 5.7)'),
]
for i, r in enumerate(rows):
    dr = data_row(tbl, *r, shade=(i % 2 == 0))
    if r[3] == 'NEW':
        set_cell_bg(dr.cells[3], 'C6EFCE')
    elif r[3] == 'REMOVED':
        set_cell_bg(dr.cells[3], 'F8D7DA')
    elif r[3] == 'Repurposed':
        set_cell_bg(dr.cells[3], 'FFF2CC')
doc.add_paragraph()

info_box(doc,
    'Pages 50684 and 50685 were already implemented in AL as part of the previous '
    '(Card-based) design iteration. Reworking to this design means: delete page 50685 '
    'entirely, strip the New/Edit/CardPageId wiring from page 50684 (leaving it as a '
    'plain read-only list), and move the validation logic currently invoked from page '
    '50685''s OnInsertRecord/OnModifyRecord into the new Post routine on codeunit 50686 '
    'instead (Section 5.9). This is a rework, not a from-scratch build - the codeunit''s '
    'core rules (Sections 5.1-5.5) do not change.',
    label='Migration Note', color='FFF2CC')

# ── 5. VALIDATION & BUSINESS LOGIC ────────────────────────────────────────────

h(doc, '5. Validation & Business Logic - Codeunit 50686 "Resource Absence Mgt."')
p(doc,
  'A single codeunit owns every Absence business rule below, mirroring the existing '
  '"Resource Handler" (codeunit 50600) pattern already used in this extension for '
  'testable, centralized resource logic. Sections 5.1-5.6 are UNCHANGED from the prior '
  'design - they describe how a Res. Capacity Entry Absence row gets built and '
  'validated, regardless of what calls into the codeunit. What changes with this '
  'revision is WHO calls it: previously the Absence Card''s OnInsertRecord/'
  'OnModifyRecord triggers called straight into these rules; now the new worksheet''s '
  'Post action does (Section 5.9), running the same rules once per unposted line.')

h(doc, '5.1  Rule: Existing-Capacity Guard', 2)
p(doc,
  'It would not be logical to record an absence without corresponding capacity. On '
  'insert/validate of an Absence row, the codeunit looks up table 160 for '
  '"Resource No." + Date + Type = Capacity + Capacity > 0.')
tbl = add_table(doc, [3, 14], 'Outcome', 'Behavior')
guard_rows = [
    ('Match found', 'The matching row''s "Duplicate Id" is copied onto the new Absence row (Section 5.3); insert proceeds.'),
    ('No match', 'Error: "Resource %1 has no capacity recorded for %2. Register capacity before recording an absence."'),
]
for i, r in enumerate(guard_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '5.2  Rule: Sign Convention', 2)
p(doc,
  'Users think in positive hours ("I was absent for 4 hours"), not negative capacity. '
  'The Absence Card exposes a positive "Hours" field; the codeunit validates Hours > 0 '
  'and writes -Hours into the underlying "Capacity" field on insert. The negative sign '
  'is an implementation detail the user never sees directly.')

h(doc, '5.3  Rule: "Duplicate Id" Inheritance', 2)
p(doc,
  'Query 50604 "Capacity Per Day Per Resource" (existing, unchanged) sums Capacity '
  'grouped by Date, Resource No. AND Duplicate Id - it does not filter on Type at all, '
  'which is exactly why this design needs no query changes. But that also means an '
  'Absence row MUST carry the same "Duplicate Id" as the Capacity row it offsets, or it '
  'nets against the wrong group (e.g. a resource with two split-shift Capacity rows for '
  'the same day, each with a different Duplicate Id). The codeunit therefore copies '
  '"Duplicate Id" from the matched Capacity row (Section 5.1), never leaves it at the '
  'field''s bare default.')

h(doc, '5.4  Rule: Overshoot Guard (recommended default - confirm in Section 9)', 2)
p(doc,
  'The codeunit blocks an Absence whose Hours exceed the linked Capacity row''s Capacity '
  '- i.e. a single absence line cannot drive that entry''s net below zero. This is a '
  'recommendation, not yet confirmed with the business - see open question 9.1.')

h(doc, '5.5  Rule: Delete Protection', 2)
p(doc,
  'Deleting a Capacity row is blocked while an Absence row still exists for the same '
  '"Resource No." + Date + "Duplicate Id" (Type = Absence) - the same lookup the '
  'existing-capacity guard (Section 5.1) already performs, just run in reverse before a '
  'delete. This prevents an Absence entry from being left behind with no corresponding '
  'Capacity for its date.')

h(doc, '5.6  Backward Compatibility', 2)
info_box(doc,
    '"Type" defaults to Capacity (enum value 0), so every existing INSERT path - the '
    'demo data codeunit and any other current writer of table 160 - continues to work '
    'completely unchanged and needs no code changes.',
    label='Compatibility', color='C6EFCE')

h(doc, '5.7  Required Fix: "Update Capacity" Wizard Must Filter by Type', 2)
p(doc,
  'Unlike a plain insert, page 50627 "Resource Capacity Settings Opt" (this extension''s '
  '"Update Capacity" action on the Resource Capacity Settings card - the actual capacity '
  'generator used day to day, a customized copy of base page 6013) RECALCULATES capacity '
  'by summing existing rows and inserting a delta. Its current logic '
  '(src/page/page_6013_ResourceCapacitySettings.al, "Update Capacity" action):')
add_code(doc,
r"""ResCapacityEntry.SetRange("Resource No.", Rec."No.");
ResCapacityEntry.SetRange(Date, TempDate);
ResCapacityEntry.SetRange("Duplicate Id", LoopCounter);
ResCapacityEntry.CalcSums(Capacity);              // <- no Type filter today
TempCapacity := ResCapacityEntry.Capacity;

ResCapacityEntry2.Capacity := NewCapacity - TempCapacity;   // adjustment row
ResCapacityEntry2.Insert(true);""")
info_box(doc,
    'Without a fix, this is a real bug, not a theoretical one: if a resource has an '
    'Absence row for a date (Capacity = -4, say) and the wizard is re-run for that date '
    'with NewCapacity = 8, CalcSums(Capacity) reads the existing sum as 8 + (-4) = 4, so '
    'the wizard inserts an adjustment of 8 - 4 = 4 - silently cancelling the absence back '
    'to a net 8, with no error and no warning to the user. REQUIRED FIX: add '
    'SetRange(Type, "Res. Capacity Entry Type"::Capacity) alongside the existing '
    '"Duplicate Id" filter, on both CalcSums(Capacity) call sites in that action '
    '(single-duplicate and multi-duplicate branches). This is the one existing object '
    'that genuinely needs a behavioral code change for this feature - not just an '
    'additive UI change - and is called out separately in Section 8.2.',
    label='Risk Found During Design Review', color='F8D7DA')

h(doc, '5.8  Post-Time Mandatory Field Validation (Worksheet Side)', 2)
p(doc,
  'Before the codeunit even attempts to build a table 160 row, the new Post routine '
  '(Section 5.9) checks each "Register Absence" line has all five of the fields the '
  'reviewer asked to be mandatory at posting - failing fast with a specific field name '
  'rather than falling through to a more generic downstream error:')
tbl = add_table(doc, [4, 13], 'Field', 'Check')
mandatory_rows = [
    ('Resource No.', 'TestField - must not be blank.'),
    ('Date', 'TestField - must not be 0D.'),
    ('Existing Capacity', 'CalcFields, then must be > 0 - the worksheet-visible surface of the existing-capacity guard (Section 5.1). Since this field nets Capacity against any already-posted Absence (Section 3.5), the check also correctly fires when a resource''s capacity for that date has already been fully used up by other posted absences, not only when it never existed.'),
    ('Hours', 'TestField - must not be 0 (and, per Section 5.2, must be positive).'),
    ('Absence Reason Code', 'TestField - must not be blank.'),
]
for i, r in enumerate(mandatory_rows):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()
info_box(doc,
    'This duplicates part of what ValidateAndPrepareAbsence (Section 5.1/5.2) already '
    'checks internally - that is intentional, not redundant cleanup material. The '
    'worksheet-side TestFields give a fast, specific "which field on which line" error '
    'before any database write is attempted; ValidateAndPrepareAbsence remains the single '
    'authoritative implementation of the actual business rules (existing-capacity match, '
    'overshoot guard, Duplicate Id inheritance) and is called unchanged from the Post loop.',
    label='Why the overlap is intentional', color='D6E4F7')

h(doc, '5.9  Posting Behavior (New)', 2)
p(doc, 'Procedure on codeunit 50686, called from the worksheet''s Post action:')
add_code(doc,
r"""procedure PostRegisterAbsence(var RegisterAbsence: Record "Register Absence")
var
    ResCapacityEntry: Record "Res. Capacity Entry";
    LastResCapacityEntry: Record "Res. Capacity Entry";
    PostedCount: Integer;
begin
    // Deliberately NOT Reset() - Post only processes whatever filter is CURRENTLY active
    // on the page (e.g. the calling Resource Card's resource). Table 160 has no
    // AutoIncrement/Number Series, so every writer assigns "Entry No." manually via
    // FindLast()+1 (matching page 50627's own convention) - this is a shared table, and
    // Reset()-ing here would let Post reach past the user's own filtered view into
    // unrelated rows sitting in the same table (a real bug found during implementation).
    if RegisterAbsence.FindSet() then
        repeat
            // 5.8 - fast, field-specific mandatory checks
            RegisterAbsence.TestField("Resource No.");
            RegisterAbsence.TestField(Date);
            RegisterAbsence.TestField("Absence Reason Code");
            RegisterAbsence.TestField(Hours);
            RegisterAbsence.CalcFields("Existing Capacity");
            if RegisterAbsence."Existing Capacity" <= 0 then
                Error(NoExistingCapacityErr, RegisterAbsence."Resource No.", RegisterAbsence.Date);

            // Build the posted ledger row and reuse the UNCHANGED validation engine
            ResCapacityEntry.Init();
            ResCapacityEntry."Resource No." := RegisterAbsence."Resource No.";
            ResCapacityEntry.Date := RegisterAbsence.Date;
            ResCapacityEntry."Absence Reason Code" := RegisterAbsence."Absence Reason Code";
            ValidateAndPrepareAbsence(ResCapacityEntry, RegisterAbsence.Hours);   // Section 5.1-5.4

            LastResCapacityEntry.Reset();
            if LastResCapacityEntry.FindLast() then
                ResCapacityEntry."Entry No." := LastResCapacityEntry."Entry No." + 1
            else
                ResCapacityEntry."Entry No." := 1;

            ResCapacityEntry.Insert(true);
            RegisterAbsence.Delete();   // posted lines do not linger in the worksheet
            PostedCount += 1;
        until RegisterAbsence.Next() = 0;

    if PostedCount = 1 then
        Message(PostedOneMsg)
    else
        if PostedCount > 1 then
            Message(PostedManyMsg, PostedCount);
end;""")
info_box(doc,
    'No Commit() is issued anywhere in this loop, so a runtime Error() on any line - a '
    'missing field, no matching capacity, an overshoot - rolls back the ENTIRE Post run '
    'via BC''s ambient transaction. Either every line in the worksheet posts and is '
    'removed, or none of them do and the worksheet is left exactly as the user had it. '
    'This all-or-nothing behavior is a natural consequence of not adding a Commit() '
    'defensively, not extra code written to enforce it - flagged in Section 9 for '
    'confirmation in case partial/best-effort posting is actually wanted instead.',
    label='All-or-Nothing Posting', color='D6E4F7')

# ── 6. PAGE & ACTION DESIGN ───────────────────────────────────────────────────

h(doc, '6. Page & Action Design')

h(doc, '6.1  Resource Card - "Absence" Action (Repurposed)', 2)
p(doc,
  'Stays on page extension 50605 "ResourceCard Opti" (existing), grouped near the '
  'current "Day Plannings (Visual)" action - only its ToolTip changes, since it now '
  'opens a read-only history view rather than an entry point. Registering a NEW '
  'absence no longer happens from the Resource Card at all; that is now exclusively '
  'the worksheet''s job (Section 6.4).')
add_code(doc,
r"""action("Absence")
{
    ApplicationArea = All;
    Caption = 'Absence';
    Image = Absence;
    RunObject = page "Resource Absence List";
    RunPageLink = "Resource No." = field("No."), Type = const(Absence);
    ToolTip = 'View this resource''s posted absence history.';
}""")

h(doc, '6.2  Page 50684 "Resource Absence List" (Repurposed - Read-Only History)', 2)
p(doc,
  'Standard list page, SourceTable "Res. Capacity Entry", SourceTableView filtered to '
  'Type = Absence. Unlike the prior design iteration, this page now has NO CardPageId, '
  'NO New action, and InsertAllowed/DeleteAllowed = false - it is purely a posted-'
  'history view. Columns: Date, Hours (shown positive via a display method or the '
  'negated Capacity value), Absence Reason Code. There is nothing left to edit here; '
  'correcting a posted absence means posting an offsetting entry via the worksheet, not '
  'modifying this list in place (consistent with how a posted journal entry is never '
  'edited directly).')
p(doc,
  'Resolves open question 9.3 (Reachability): the worksheet is reached from THIS list, '
  'not from the Resource Card directly, via a plain "Register Absence" action that '
  'opens page 50688 with no RunPageLink/pre-filter - the user picks the Resource No. '
  'for each line on the worksheet itself.')
add_code(doc,
r"""action("RegisterAbsence")
{
    ApplicationArea = All;
    Caption = 'Register Absence';
    Image = Register;
    RunObject = page "Register Absence";
}""")

h(doc, '6.3  Page 50685 "Resource Absence Card" - REMOVED', 2)
info_box(doc,
    'This page existed in the prior design iteration as the direct entry point into '
    'table 160. It is dropped entirely by this revision: all new-entry data capture '
    'moves to the "Register Absence" worksheet (Section 6.4), which posts into table 160 '
    'instead of a Card writing to it directly. If already implemented, delete this page '
    'object as part of reworking to this design (Section 4 Migration Note).',
    label='Removed from Design', color='F8D7DA')

h(doc, '6.4  Page 50688 "Register Absence" [NEW] - Worksheet', 2)
p(doc,
  'Journal-style worksheet, PageType = Worksheet (or a fully editable List, if a plain '
  'list better matches this extension''s existing UI conventions), SourceTable '
  '"Register Absence" (Section 3.5), Editable/InsertAllowed/ModifyAllowed/DeleteAllowed '
  '= true. Lets the user key in several absence lines - across different resources and '
  'dates - before posting any of them.')
tbl = add_table(doc, [4.5, 4.5, 8], 'Column', 'Editable?', 'Notes')
worksheet_fields = [
    ('Resource No.', 'Yes', 'TableRelation to Resource.'),
    ('Resource Name', 'No', 'Plain page-level lookup (Resource.Get() in OnAfterGetRecord / a page variable) - not a stored table field, matching the same non-FlowField convention used throughout this design.'),
    ('Date', 'Yes', 'The absence date for this line.'),
    ('Absence Reason Code', 'Yes', 'Lookup to table 5206 "Cause of Absence" / page "Causes of Absence".'),
    ('Hours', 'Yes', 'Positive entry. No pre-fill default - user always types it (per the earlier-confirmed design decision, carried over unchanged from the Card iteration).'),
    ('Existing Capacity', 'No (FlowField)', 'Read-only, visible before posting - the worksheet-side surface of the existing-capacity guard (Section 5.8).'),
]
for i, r in enumerate(worksheet_fields):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()
p(doc, 'Actions:')
add_code(doc,
r"""action(Post)
{
    ApplicationArea = All;
    Caption = 'Post';
    Image = Post;
    InFooterBar = true;
    ToolTip = 'Validate and post all absence lines on this worksheet into Res. Capacity Entry.';

    trigger OnAction()
    var
        ResourceAbsenceMgt: Codeunit "Resource Absence Mgt.";
    begin
        ResourceAbsenceMgt.PostRegisterAbsence(Rec);   // Section 5.9
    end;
}""")

# ── 7. DATA FLOW ──────────────────────────────────────────────────────────────

h(doc, '7. Data Flow')
p(doc, '7.1  Entry + Posting (new, replaces the old Card flow):')
add_code(doc,
r""""Register Absence" worksheet  (Page 50688, table 50687 - unposted)
      |
      |  user keys in N lines: Resource No., Date, Absence Reason Code, Hours
      |  "Existing Capacity" FlowField shows live, before posting, the NET
      |  remaining capacity for each line (Capacity minus already-posted Absence)
      v
Post action  ->  Codeunit 50686 "Resource Absence Mgt.".PostRegisterAbsence  (Section 5.9)
      |
      |  for each line CURRENTLY VISIBLE on the page (no Reset - a shared table,
      |  Post never reaches past the page's own filter), in order:
      |-- 5.8  TestField Resource No. / Date / Absence Reason Code / Hours;
      |          CalcFields+check Existing Capacity > 0 (net remaining, not just "exists")
      |-- 5.1  find matching Capacity row (Resource No. + Date + Type=Capacity)
      |          not found -> Error, ENTIRE POST ROLLS BACK (no Commit in the loop)
      |          found     -> capture its Duplicate Id
      |-- 5.4  Hours <= matched row's Capacity ?  (overshoot guard)
      |-- 5.2  negate Hours -> Capacity field
      |-- INSERT into Table 160 "Res. Capacity Entry"
      |       (Type = Absence, Capacity = -Hours, Duplicate Id = copied)
      |-- DELETE the now-posted "Register Absence" line
      v
Existing, UNCHANGED aggregation automatically nets the result:
   - Query 50604 "Capacity Per Day Per Resource"   (SUM Capacity by Date/Resource No./Duplicate Id)
   - Any other existing SUM/FlowField reader of "Res. Capacity Entry".Capacity""")
p(doc, '7.2  History Viewing + Worksheet Entry Point:')
add_code(doc,
r"""Resource Card
      |
      |  action "Absence"  (RunPageLink: Resource No. + Type = Absence)
      v
Resource Absence List  (Page 50684, read-only, filtered to Type = Absence)
      |
      |  action "Register Absence"  (no RunPageLink - opens unfiltered)
      v
"Register Absence" worksheet  (Page 50688 - see 7.1 for the posting flow)""")

# ── 8. IMPACTED OBJECTS SUMMARY ───────────────────────────────────────────────

h(doc, '8. Impacted Objects Summary')
p(doc, 'For a reviewer to gauge blast radius at a glance:')

h(doc, '8.1  Touched but NOT code-changed (confirmed by design)', 2)
tbl = add_table(doc, [5, 12], 'Object', 'Why it needs no change')
untouched = [
    ('Query 50604 "Capacity Per Day Per Resource"', 'Sums Capacity with no Type filter - already nets Capacity/Absence automatically (Section 5.3). Read-only aggregation, so unlike page 50627 there is no "recompute and re-insert a delta" step that could be fooled.'),
    ('Codeunit 50602 "CreateDemoData"', 'Plain Insert of Capacity rows - new Type field defaults to Capacity, so demo data generation is unaffected (Section 5.6).'),
]
for i, r in enumerate(untouched):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '8.2  New', 2)
tbl = add_table(doc, [5, 12], 'Object', 'Purpose')
new_objs = [
    ('Table 50687 "Register Absence"', 'Unposted worksheet staging table (Section 3.5).'),
    ('Page 50688 "Register Absence"', 'Worksheet page with the Post action (Section 6.4).'),
]
for i, r in enumerate(new_objs):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '8.3  Modified', 2)
tbl = add_table(doc, [5, 12], 'Object', 'Change')
modified = [
    ('Table Extension 50606 "ResCapacityEntry Opt"', '+ 2 fields (Section 3.1) - unchanged from prior iteration.'),
    ('Codeunit 50686 "Resource Absence Mgt."', 'Core rules (Sections 5.1-5.6) unchanged; + new PostRegisterAbsence entry point (Section 5.9) called from the worksheet instead of a Card''s triggers.'),
    ('Page 50684 "Resource Absence List"', 'Repurposed to a plain read-only history list - New/Edit/CardPageId wiring removed (Section 6.2).'),
    ('Page Extension 50605 "ResourceCard Opti"', '"Absence" action retained, ToolTip updated to reflect its now-read-only purpose (Section 6.1).'),
    ('Page 50629 "Res. Capacity FactBox Part"', '+ "Type" column, so the factbox on the Resource Card no longer shows an unexplained second row per date once Absence entries exist.'),
    ('Page 50627 "Resource Capacity Settings Opt"', 'BEHAVIORAL fix, not additive: both CalcSums(Capacity) call sites in the "Update Capacity" action need a Type = Capacity filter added, or re-running that wizard silently cancels out recorded absences (Section 5.7).'),
]
for i, r in enumerate(modified):
    data_row(tbl, *r, shade=(i % 2 == 0))
doc.add_paragraph()

h(doc, '8.4  Removed', 2)
tbl = add_table(doc, [5, 12], 'Object', 'Reason')
removed_objs = [
    ('Page 50685 "Resource Absence Card"', 'Superseded by the worksheet - direct Card entry into table 160 is no longer part of the design (Section 6.3).'),
]
for i, r in enumerate(removed_objs):
    data_row(tbl, *r, shade=(i % 2 == 0))

# ── 9. OPEN QUESTIONS ──────────────────────────────────────────────────────────

h(doc, '9. Open Questions / Assumptions for Reviewer')
p(doc,
  'The three business-rule questions raised in the prior (Card-based) design iteration '
  'were already confirmed with the reviewer and carry over unchanged to this revision: '
  'overshoot is blocked (Section 5.4), multiple absence lines per Resource/Date are '
  'allowed (Section 5.3), and Hours never pre-fills a default (Section 6.4). '
  'Reachability (formerly 9.3) is now also confirmed: a plain "Register Absence" '
  'action on the read-only history list (Page 50684, Section 6.2) opens the worksheet '
  'unfiltered - not an action on the Resource Card. The two questions below remain '
  'open, specific to the posting mechanism itself.')
info_box(doc,
    '9.1  All-or-nothing posting (Section 5.9) - if one line among several on the '
    'worksheet fails validation, should the ENTIRE Post action roll back (current '
    'design, the natural default with no mid-loop Commit), or should valid lines post '
    'while invalid ones are skipped/left behind with an error summary?\n\n'
    '9.2  Batch/template concept - should "Register Absence" support multiple named '
    'batches like General Journal Batches (so different users/teams keep separate '
    'worksheets), or is a single shared worksheet list sufficient (current design, '
    'simpler)?',
    label='Confirm before implementation', color='FFF2CC')

# ── 10. OBJECT ID QUICK REFERENCE ─────────────────────────────────────────────

h(doc, '10. Object ID Quick Reference')
tbl = add_table(doc, [3, 3, 11], 'Type', 'ID', 'Name')
quick_ref = [
    ('Enum', '50682', 'Res. Capacity Entry Type'),
    ('Table Extension', '50606', 'ResCapacityEntry Opt (extended)'),
    ('Table', '50687', 'Register Absence (NEW - unposted worksheet staging)'),
    ('Page (Worksheet)', '50688', 'Register Absence (NEW - Post action)'),
    ('Codeunit', '50686', 'Resource Absence Mgt. (+ PostRegisterAbsence)'),
    ('Page (List)', '50684', 'Resource Absence List (repurposed - read-only history)'),
    ('Page (Card)', '50685', 'Resource Absence Card (REMOVED)'),
    ('Page Extension', '50605', 'ResourceCard Opti (extended)'),
    ('Page', '50629', 'Res. Capacity FactBox Part (extended)'),
    ('Page', '50627', 'Resource Capacity Settings Opt (behavioral fix - Section 5.7)'),
    ('Table (reference, unchanged)', '160', 'Res. Capacity Entry (base app)'),
    ('Table (reference, unchanged)', '5206', 'Cause of Absence (base app)'),
]
for i, r in enumerate(quick_ref):
    dr = data_row(tbl, *r, shade=(i % 2 == 0))
    if r[1] == '50627':
        set_cell_bg(dr.cells[1], 'F8D7DA')
    elif r[1] == '50685':
        set_cell_bg(dr.cells[1], 'F8D7DA')
    elif r[1] in ('50687', '50688'):
        set_cell_bg(dr.cells[1], 'C6EFCE')

# ── SAVE ───────────────────────────────────────────────────────────────────────

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Resource Absence Management.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
