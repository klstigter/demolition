"""
generate_capacity_vs_requested_chart_doc.py
Generates the user/explainer guide for the stacked "Capacity vs Requested" chart on
page 50692 "Requested vs Capacity Skl Dhx" as a Word (.docx) file using python-docx.

Audience: Ahmad (technical, but wants a plain-language explanation he can later re-explain
to end users) - not a from-scratch AL design doc like the other generators in this folder.

Run:  python generate_capacity_vs_requested_chart_doc.py
Output: Capacity vs Requested Chart - User Guide.docx  (same folder as this script)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ─── helpers (same conventions as the other docs/generate_*.py scripts) ─────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def header_row(table, *texts, bg='1F3964'):
    row = table.rows[0]
    for i, text in enumerate(texts):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)


def data_row(table, *values, shade=False):
    row = table.add_row()
    for i, val in enumerate(values):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        if shade:
            set_cell_bg(cell, 'EBF3FB')
        p = cell.paragraphs[0]
        p.add_run(str(val)).font.size = Pt(10)
    return row


def add_table(doc, col_widths_cm, *header_texts, bg='1F3964'):
    cols = len(header_texts)
    tbl = doc.add_table(rows=1, cols=cols)
    tbl.style = 'Table Grid'
    header_row(tbl, *header_texts, bg=bg)
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[j].width = Cm(w)
    return tbl


def add_swatch_table(doc, col_widths_cm, header_texts, rows_with_hex):
    """Like add_table, but the first data column is filled with the row's own colour swatch."""
    tbl = add_table(doc, col_widths_cm, *header_texts)
    for i, (hexval, *rest) in enumerate(rows_with_hex):
        dr = data_row(tbl, hexval, *rest, shade=(i % 2 == 0))
        set_cell_bg(dr.cells[0], hexval.lstrip('#'))
        # white text on dark swatches so the hex code stays legible
        for run in dr.cells[0].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
    return tbl


def add_code(doc, code: str):
    """Monospace code/formula block with light-blue background."""
    for line in code.split('\n'):
        para = doc.add_paragraph()
        para.style = doc.styles['No Spacing']
        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EDF2F8')
        pPr.append(shd)
    doc.add_paragraph()


def info_box(doc, text: str, label='Note', color='D6E4F7'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    run = p.add_run(f'{label}\n{text}')
    run.font.size = Pt(10)
    doc.add_paragraph()


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text=''):
    return doc.add_paragraph(text)


def bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


# ════════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Capacity vs Requested Chart')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run('User Guide - what the chart shows, what "Scenario" means, and how to explain it to end users').font.size = Pt(14)

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(
    f'Page: 50692 "Skill Requested/Capacity"  |  Extension: DailyOptimizer\n'
    f'Generated: {datetime.datetime.now().strftime("%d %B %Y %H:%M")}'
).font.size = Pt(10)

doc.add_paragraph()
info_box(doc,
    'This is a plain-language explainer, not an AL technical design document (see docs/ for '
    'those). It is written so you (Ahmad) can review the logic once here, then explain the '
    'chart to end users in your own words without needing to re-derive the formulas each '
    'time. Section 8 is a ready-to-use "one sentence per concept" cheat sheet for that.',
    label='Purpose of this document', color='D6E4F7')

doc.add_page_break()

# ── 1. WHAT THE CHART SHOWS ───────────────────────────────────────────────────

h(doc, '1. What the Chart Shows')
p(doc,
  'The chart on page 50692 answers one question for a planner: "for each day this week, how '
  'much capacity do we have left, and how much is still being asked for, broken down by skill?" '
  'It shows Monday through Friday of the currently displayed week, and for every day there are '
  'two side-by-side bars:')
bullet(doc, 'Capacity bar - how much resource capacity exists that day, split into what is already assigned vs. still free (internal vs. external).')
bullet(doc, 'Requested bar - how much work is being requested that day, split into what is already assigned vs. still open per skill (Electrician, Sanitary, ...).')
p(doc,
  'Both bars are "stacked" - each one is built by piling coloured segments on top of each '
  'other, so the total height of the bar is the sum of its segments, and you can read both '
  'the total and the breakdown at a glance.')

# ── 2. WHERE IT CAME FROM ─────────────────────────────────────────────────────

h(doc, '2. Where It Came From')
p(doc,
  'This was originally prototyped in Excel ("Cappacity vs Requested.xlsx") with hand-typed '
  'example numbers and a "Scenario" cell that swapped between 5 pre-built sets of numbers. '
  'The BC version keeps the exact same visual idea (two stacked bars per day, a Scenario '
  'control) but replaces the hand-typed numbers with LIVE numbers computed from real Day '
  'Planning, Resource Capacity, Resource, and Skill Code data - see Section 4 for exactly '
  'how each segment is calculated.')

# ── 3. READING THE COLOURS ────────────────────────────────────────────────────

h(doc, '3. Reading the Colours')
p(doc, 'Every segment has a fixed colour, so the same colour always means the same thing across every day and every time you open the page:')
add_swatch_table(doc, [3, 4, 10], ['Colour', 'Segment', 'Appears on'],
    [('#548235', 'Assigned', 'Both bars (identical value on both - see Section 4.1)'),
     ('#8EA9DB', 'Internal', 'Capacity bar only'),
     ('#B4C6E7', 'External (red outline)', 'Capacity bar only'),
     ('#C55A11', 'Skill 1 (darkest)', 'Requested bar only'),
     ('#ED7D31', 'Skill 2', 'Requested bar only'),
     ('#F4B183', 'Skill 3', 'Requested bar only'),
     ('#F8CBAD', 'Skill 4', 'Requested bar only'),
     ('#FBE5D6', 'Skill 5 (lightest, then repeats)', 'Requested bar only')])
p(doc,
  'The skill colours are assigned in the order skills are found for that week, cycling back '
  'to the first colour if there are more than 5 skills with open demand that week. A skill '
  'only gets a colour/segment at all if it actually has open (unassigned) demand somewhere '
  'in the displayed week - a skill nobody is asking for that week simply does not appear, '
  'instead of showing an empty zero-height segment.')

# ── 4. WHAT EACH SEGMENT MEANS ────────────────────────────────────────────────

h(doc, '4. What Each Segment Actually Means')

h(doc, '4.1  Assigned (green) - the shared baseline', 2)
p(doc,
  'This is simply the total hours already assigned to someone for that day. It is the ONE '
  'number that appears identically on both bars (Capacity and Requested) - because it is the '
  'same underlying fact viewed two ways: "capacity already used up" and "requested work '
  'already spoken for". Everything else stacks on top of this shared base.')
add_code(doc, 'Assigned = SUM(Day Planning."Assigned Hours") for that day')

h(doc, '4.2  Internal / External (blue tones) - free capacity, Capacity bar only', 2)
p(doc,
  'What is left over after Assigned is taken out of total capacity, split by whether the '
  'resource is an internal employee or an external/subcontracted one (Resource."Is External").')
add_code(doc,
r"""Internal (free) = SUM(Capacity, internal resources)  -  SUM(Assigned Hours, internal resources)
External (free) = SUM(Capacity, external resources)  -  SUM(Assigned Hours, external resources)""")
info_box(doc,
    'These can legitimately go negative. If a day has been over-assigned relative to its '
    'recorded capacity, the segment shows negative on purpose - that is a real, useful signal '
    '("we have committed more than we actually have capacity for"), not a bug to hide.',
    label='Why a segment can be negative', color='FFF2CC')

h(doc, '4.3  Per-skill segments (orange tones) - open demand, Requested bar only', 2)
p(doc,
  'For each skill, this is the hours requested for that day that are NOT yet assigned to '
  'anyone. Already-assigned work for that skill is already counted inside "Assigned" '
  '(Section 4.1), so it is deliberately excluded here to avoid counting it twice.')
add_code(doc, 'Skill segment = SUM(Requested Hours) for that skill, only lines with no Assigned Resource yet')

# ── 5. THE SCENARIO CONTROL ───────────────────────────────────────────────────

h(doc, '5. The Scenario Control')
p(doc,
  'The "Scenario" field is a number from 0 to 5 that simulates "how far into the week are '
  'we". It closes days one at a time, starting from Monday: a closed day collapses down to '
  'just its green Assigned segment (everything else on that day goes to zero), because a day '
  'that is already done no longer needs a detailed breakdown - you only need to see what was '
  'actually assigned.')

add_table(doc, [2, 3.5, 10.5], 'Scenario', 'Name', 'Meaning')
scenario_rows = [
    ('0', 'Open', 'Nothing closed yet - every day (Mon-Fri) shows its full breakdown.'),
    ('1', 'Mon Closed', 'Monday collapsed to Assigned-only. Tue-Fri still fully open.'),
    ('2', 'Tue Closed', 'Monday + Tuesday collapsed. Wed-Fri still fully open.'),
    ('3', 'Wed Closed', 'Monday-Wednesday collapsed. Thu-Fri still fully open.'),
    ('4', 'Thu Closed', 'Monday-Thursday collapsed. Only Friday still fully open.'),
    ('5', 'Closed', 'The whole week is collapsed to Assigned-only.'),
]
tbl = doc.tables[-1]
for i, row in enumerate(scenario_rows):
    data_row(tbl, *row, shade=(i % 2 == 0))
doc.add_paragraph()

p(doc,
  'The rule behind the table: weekday i (Monday = 1 ... Friday = 5) is closed whenever '
  'i <= Scenario. So the name is always just "the last closed weekday" - once you know that '
  'one rule, you can work out any Scenario value without memorising the table.')
info_box(doc,
    'When the page opens (or you move to a different week with Previous/Today/Next), Scenario '
    'defaults automatically to how many weekdays of that week are already in the past relative '
    'to today\'s date - so a Tuesday viewing of the current week opens on Scenario 1 (Mon '
    'Closed) by itself. You can still change it manually to "play" with other scenarios, '
    'exactly like the original Excel prototype\'s B18 cell.',
    label='Where the starting value comes from', color='D6E4F7')

# ── 6. WORKED EXAMPLE ─────────────────────────────────────────────────────────

h(doc, '6. Worked Example')
p(doc,
  'A small concrete example (the same numbers used in the automated test suite, so they are '
  'independently verified - see Section 7). One internal resource has 10 hours of Capacity '
  'and 6 Assigned Hours on Monday, Tuesday, and Wednesday, plus 3 hours of open (unassigned) '
  'demand for one skill each of those days. Scenario is set to 2 (Tue Closed).')

add_table(doc, [3, 3, 3, 3, 3.5], 'Day', 'Assigned', 'Internal (free)', 'Skill segment', 'State')
example_rows = [
    ('Monday (i=1)', '6', '0  (collapsed)', '0  (collapsed)', 'Closed (1 <= 2)'),
    ('Tuesday (i=2)', '6', '0  (collapsed)', '0  (collapsed)', 'Closed (2 <= 2)'),
    ('Wednesday (i=3)', '6', '4  (10 - 6)', '3', 'Open (3 > 2)'),
]
tbl = doc.tables[-1]
for i, row in enumerate(example_rows):
    data_row(tbl, *row, shade=(i % 2 == 0))
doc.add_paragraph()
p(doc,
  'Notice Assigned stays 6 on every day regardless of Scenario - only the Internal/skill '
  'segments get zeroed on closed days. Wednesday is the first open day, so it shows its real '
  'numbers: 4 hours of free internal capacity (10 capacity minus 6 already assigned) and 3 '
  'hours of open demand for that skill.')

# ── 7. HOW THIS WAS VERIFIED ──────────────────────────────────────────────────

h(doc, '7. How This Was Verified')
p(doc, 'Two layers of checking back this feature, so the numbers on the chart can be trusted:')
bullet(doc, 'Automated regression tests - codeunit 60024 "Skill Capacity Chart Tests" (test\\SkillCapacityChart.Test.Codeunit.al), 9 tests covering: Assigned being identical on both bars, the Internal/External split following Resource."Is External", already-assigned lines being excluded from skill segments, a skill with zero open demand not appearing at all, and the Scenario collapse boundary for every value 1 through 4.')
bullet(doc, 'Live data sanity check - the Internal/External classification and Assigned/Requested field mapping were cross-checked against real records in the NL_Test sandbox (e.g. resource DRM00003 has no Vendor No. and is Internal; resource DRP002 has Vendor No. DRV002 and is External; already-assigned Day Planning lines have Assigned Hours equal to Requested Hours, matching the "Assigned" concept).')
info_box(doc,
    'All 9 automated tests currently pass. Re-run them any time via the AL test tooling on '
    'codeunit 60024 to confirm the calculations still hold after any future change.',
    label='Current status', color='C6EFCE')

# ── 8. CHEAT SHEET FOR EXPLAINING TO END USERS ────────────────────────────────

h(doc, '8. Cheat Sheet - Explaining This to End Users')
p(doc, 'One sentence per concept, in plain language, no field names or formulas:')
tbl = add_table(doc, [5, 12], 'When they ask about...', 'Say this')
cheat_rows = [
    ('The chart in general', '"For each day this week, it shows how much capacity we have left and how much people are still asking for, split by skill."'),
    ('The green segment', '"That\'s the work already assigned - it\'s the same number on both bars because it\'s the same hours, just shown from two angles."'),
    ('The blue segments', '"That\'s how much capacity is still free - split into our own people (Internal) and outside/subcontracted (External, the one with the red outline)."'),
    ('The orange segments', '"Each colour is a different skill still needed that isn\'t assigned to anyone yet - electricians, sanitary, and so on."'),
    ('A negative blue segment', '"That means we\'ve already assigned more than we actually have capacity for that day - worth a closer look."'),
    ('The Scenario field', '"It lets you simulate how far into the week we are - closing off days one at a time from Monday, so you can see just the days still open for planning."'),
    ('Why a day looks empty', '"That day is \'closed\' under the current Scenario - it\'s already committed, so we only show what was assigned, not a full breakdown."'),
]
for i, row in enumerate(cheat_rows):
    data_row(tbl, *row, shade=(i % 2 == 0))

# ── SAVE ───────────────────────────────────────────────────────────────────────

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Capacity vs Requested Chart - User Guide.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
